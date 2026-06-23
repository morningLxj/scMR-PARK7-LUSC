from __future__ import annotations

import gzip
import math
import os
import re
import shutil
import tarfile
import warnings
import zipfile
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from PIL import Image, ImageDraw
from scipy import sparse, stats
from scipy.io import mmread


REPO_ROOT = Path(__file__).resolve().parents[2]
ROOT = Path(os.environ.get("SCMR_PROJECT_ROOT", REPO_ROOT))
WORK = Path(os.environ.get("SCMR_WORK_DIR", ROOT / "work" / "final_unfinished_analyses_20260612"))
SC_DIR = WORK / "public_scRNA_GSE200972"
OUT = WORK / "outputs"
PLOT = OUT / "figures"
TABLE = OUT / "tables"

FRONTIERS_ROOT = Path(os.environ.get("SCMR_SUBMISSION_PACKAGE", ROOT / "submission_package"))
FRONTIERS_68 = Path(os.environ.get("SCMR_LEGACY_SUBMISSION_PACKAGE", ROOT / "legacy_submission_package"))

for p in [WORK, OUT, PLOT, TABLE]:
    p.mkdir(parents=True, exist_ok=True)


def bh_fdr(pvals: pd.Series) -> pd.Series:
    p = pvals.astype(float).to_numpy()
    out = np.full_like(p, np.nan, dtype=float)
    ok = np.isfinite(p)
    if ok.sum() == 0:
        return pd.Series(out, index=pvals.index)
    pv = p[ok]
    order = np.argsort(pv)
    ranked = pv[order]
    n = len(ranked)
    adj = ranked * n / np.arange(1, n + 1)
    adj = np.minimum.accumulate(adj[::-1])[::-1]
    adj = np.clip(adj, 0, 1)
    restored = np.empty_like(adj)
    restored[order] = adj
    out[ok] = restored
    return pd.Series(out, index=pvals.index)


def read_10x(prefix: str) -> tuple[sparse.csc_matrix, pd.DataFrame, pd.Series]:
    matrix_path = SC_DIR / f"{prefix}_matrix.mtx.gz"
    feature_path = SC_DIR / f"{prefix}_features.tsv.gz"
    barcode_path = SC_DIR / f"{prefix}_barcodes.tsv.gz"
    mat = mmread(str(matrix_path)).tocsc()
    features = pd.read_csv(feature_path, sep="\t", header=None)
    if features.shape[1] >= 2:
        features = features.iloc[:, :2]
        features.columns = ["gene_id", "gene_symbol"]
    else:
        features.columns = ["gene_symbol"]
        features["gene_id"] = features["gene_symbol"]
    barcodes = pd.read_csv(barcode_path, sep="\t", header=None)[0].astype(str)
    return mat, features, barcodes


MARKERS = {
    "Epithelial/tumor": ["EPCAM", "KRT8", "KRT18", "KRT19", "KRT5", "KRT14", "TP63"],
    "B cells": ["MS4A1", "CD79A", "CD79B", "CD19", "CD74", "BANK1"],
    "Plasma cells": ["MZB1", "SDC1", "JCHAIN", "IGHG1", "XBP1"],
    "T cells": ["CD3D", "CD3E", "TRAC", "CD2", "IL7R"],
    "NK cells": ["NKG7", "GNLY", "PRF1", "KLRD1"],
    "Myeloid cells": ["LYZ", "LST1", "TYROBP", "CD68", "FCGR3A"],
    "Fibroblasts": ["COL1A1", "COL1A2", "DCN", "LUM", "COL3A1"],
    "Endothelial cells": ["PECAM1", "VWF", "KDR", "CLDN5", "ENG"],
    "Mast cells": ["TPSAB1", "TPSB2", "KIT", "CPA3"],
}


REDOX_EXCLUDE = {
    "PARK7", "DJ1", "DJ-1", "NFE2L2", "NRF2", "KEAP1", "HMOX1", "NQO1",
    "GCLC", "GCLM", "GSR", "GSS", "SLC7A11", "TXN", "TXN2", "TXNRD1",
    "TXNRD2", "SRXN1", "SOD1", "SOD2", "SOD3", "CAT", "GPX1", "GPX2",
    "GPX3", "GPX4", "PRDX1", "PRDX2", "PRDX3", "PRDX4", "PRDX5",
    "PRDX6", "GLRX", "G6PD", "FTH1", "FTL",
}


def normalize_log2_cpm(mat: sparse.csc_matrix) -> sparse.csc_matrix:
    lib = np.asarray(mat.sum(axis=0)).ravel().astype(float)
    lib[lib == 0] = 1.0
    norm = mat.copy().astype(float)
    norm = norm @ sparse.diags(1e4 / lib)
    norm.data = np.log2(1.0 + norm.data)
    return norm


def gene_index(features: pd.DataFrame) -> dict[str, list[int]]:
    mapping: dict[str, list[int]] = {}
    for i, g in enumerate(features["gene_symbol"].astype(str)):
        mapping.setdefault(g.upper(), []).append(i)
    return mapping


def marker_scores(log_mat: sparse.csc_matrix, features: pd.DataFrame) -> pd.DataFrame:
    idx = gene_index(features)
    scores = {}
    for label, genes in MARKERS.items():
        rows = [j for g in genes for j in idx.get(g.upper(), [])]
        if rows:
            scores[label] = np.asarray(log_mat[rows, :].mean(axis=0)).ravel()
        else:
            scores[label] = np.zeros(log_mat.shape[1])
    return pd.DataFrame(scores)


def run_public_lusc_scrna() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    samples = {
        "GSM6047623_P1_T_R_I": "LUSC tumor P1 region I",
        "GSM6047625_P1_T_R_M": "LUSC tumor P1 region M",
    }
    cell_frames = []
    b_mats = []
    b_features = None
    b_cell_meta = []

    for prefix, label in samples.items():
        mat, features, barcodes = read_10x(prefix)
        log_mat = normalize_log2_cpm(mat)
        scores = marker_scores(log_mat, features)
        max_score = scores.max(axis=1)
        assigned = scores.idxmax(axis=1)
        assigned[max_score < 0.20] = "Low-confidence/other"

        idx = gene_index(features)
        park7_rows = idx.get("PARK7", [])
        if not park7_rows:
            raise RuntimeError(f"PARK7 was not found in {prefix}")
        park7 = np.asarray(log_mat[park7_rows, :].mean(axis=0)).ravel()

        meta = pd.DataFrame({
            "dataset": "GSE200972",
            "sample": prefix,
            "sample_label": label,
            "barcode": barcodes.to_numpy(),
            "assigned_cell_type": assigned.to_numpy(),
            "assignment_score": max_score.to_numpy(),
            "PARK7_log2CPM": park7,
            "PARK7_detected": park7 > 0,
        })
        cell_frames.append(meta)

        bmask = meta["assigned_cell_type"].isin(["B cells", "Plasma cells"]).to_numpy()
        if bmask.sum() >= 20:
            b_mats.append(mat[:, bmask])
            b_cell_meta.append(meta.loc[bmask, ["dataset", "sample", "barcode", "assigned_cell_type", "PARK7_log2CPM"]])
            b_features = features

    cell_df = pd.concat(cell_frames, ignore_index=True)
    cell_df.to_csv(TABLE / "S7_GSE200972_LUSC_scRNA_PARK7_cell_level.csv", index=False)

    summary = (
        cell_df.groupby(["sample", "sample_label", "assigned_cell_type"], as_index=False)
        .agg(
            n_cells=("barcode", "size"),
            PARK7_mean_log2CPM=("PARK7_log2CPM", "mean"),
            PARK7_median_log2CPM=("PARK7_log2CPM", "median"),
            PARK7_pct_detected=("PARK7_detected", lambda x: 100 * float(np.mean(x))),
        )
    )
    summary.to_csv(TABLE / "S7_GSE200972_LUSC_scRNA_PARK7_celltype_summary.csv", index=False)

    if not b_mats or b_features is None:
        raise RuntimeError("No B-lineage cells were assigned in the GSE200972 LUSC samples.")
    b_all = sparse.hstack(b_mats, format="csc")
    b_meta = pd.concat(b_cell_meta, ignore_index=True)
    cmap_signature = build_cmap_signature_from_b_cells(b_all, b_features, b_meta)
    make_s7_figure(cell_df, summary)
    return cell_df, summary, cmap_signature


def build_cmap_signature_from_b_cells(
    b_counts: sparse.csc_matrix, features: pd.DataFrame, b_meta: pd.DataFrame
) -> pd.DataFrame:
    log_b = normalize_log2_cpm(b_counts)
    idx = gene_index(features)
    park7_rows = idx.get("PARK7", [])
    park7 = np.asarray(log_b[park7_rows, :].mean(axis=0)).ravel()
    positive = park7[park7 > 0]
    cutoff = float(np.median(positive)) if positive.size else float(np.median(park7))
    high = park7 > cutoff
    low = ~high
    if high.sum() < 10 or low.sum() < 10:
        cutoff = float(np.median(park7))
        high = park7 > cutoff
        low = ~high

    rows = []
    symbols = features["gene_symbol"].astype(str).to_numpy()
    log_high = log_b[:, high]
    log_low = log_b[:, low]
    mean_high = np.asarray(log_high.mean(axis=1)).ravel()
    mean_low = np.asarray(log_low.mean(axis=1)).ravel()
    pct_high = np.asarray((b_counts[:, high] > 0).mean(axis=1)).ravel()
    pct_low = np.asarray((b_counts[:, low] > 0).mean(axis=1)).ravel()

    for i, gene in enumerate(symbols):
        if not gene or gene == "nan":
            continue
        rows.append({
            "gene": gene,
            "avg_log2FC": mean_high[i] - mean_low[i],
            "mean_log2CPM_PARK7_high": mean_high[i],
            "mean_log2CPM_PARK7_low": mean_low[i],
            "pct_high": pct_high[i],
            "pct_low": pct_low[i],
        })
    sig = pd.DataFrame(rows)
    sig = sig.groupby("gene", as_index=False).agg({
        "avg_log2FC": "max",
        "mean_log2CPM_PARK7_high": "max",
        "mean_log2CPM_PARK7_low": "max",
        "pct_high": "max",
        "pct_low": "max",
    })
    sig["max_pct"] = sig[["pct_high", "pct_low"]].max(axis=1)
    sig["excluded_manual_redox_or_anchor"] = sig["gene"].str.upper().isin(REDOX_EXCLUDE)
    sig["excluded_low_detection"] = sig["max_pct"] < 0.05
    sig["eligible_for_CMap"] = ~(sig["excluded_manual_redox_or_anchor"] | sig["excluded_low_detection"])
    sig["PARK7_group_cutoff_log2CPM"] = cutoff
    sig["n_B_lineage_cells"] = b_counts.shape[1]
    sig["n_PARK7_high"] = int(high.sum())
    sig["n_PARK7_low"] = int(low.sum())
    sig = sig.sort_values("avg_log2FC", ascending=False)
    sig.to_csv(TABLE / "CMap_GSE200972_LUSC_Blineage_PARK7_high_vs_low_all_genes.csv", index=False)

    eligible = sig[sig["eligible_for_CMap"]].copy()
    up = eligible.sort_values("avg_log2FC", ascending=False).head(50)
    down = eligible.sort_values("avg_log2FC", ascending=True).head(50)
    up["CMap_direction"] = "UP"
    down["CMap_direction"] = "DOWN"
    top = pd.concat([up, down], ignore_index=True)
    top.to_csv(TABLE / "SupplementaryTableS5_CMap_Top50_Filtered_Signature_Source.csv", index=False)

    (TABLE / "CMap_Input_Up_top50_filtered.txt").write_text("\n".join(up["gene"]) + "\n", encoding="utf-8")
    (TABLE / "CMap_Input_Down_top50_filtered.txt").write_text("\n".join(down["gene"]) + "\n", encoding="utf-8")
    (TABLE / "CMap_Input_Tags_top50_filtered.txt").write_text(
        "UP_TAGS:\n" + "\n".join(up["gene"]) + "\n\nDOWN_TAGS:\n" + "\n".join(down["gene"]) + "\n",
        encoding="utf-8",
    )
    b_meta.assign(PARK7_group=np.where(high, "PARK7_high", "PARK7_low")).to_csv(
        TABLE / "CMap_GSE200972_Blineage_cells_PARK7_groups.csv", index=False
    )
    return top


def make_s7_figure(cell_df: pd.DataFrame, summary: pd.DataFrame) -> None:
    order = (
        summary.groupby("assigned_cell_type")["n_cells"].sum()
        .sort_values(ascending=False).index.tolist()
    )
    colors = {
        "Epithelial/tumor": "#2f6f9f",
        "B cells": "#c74732",
        "Plasma cells": "#ef8a47",
        "T cells": "#4d8f3a",
        "NK cells": "#7b529d",
        "Myeloid cells": "#6b6b6b",
        "Fibroblasts": "#b88a2d",
        "Endothelial cells": "#288f8f",
        "Mast cells": "#a04d73",
        "Low-confidence/other": "#bdbdbd",
    }
    fig = plt.figure(figsize=(13, 10), constrained_layout=True)
    gs = fig.add_gridspec(2, 2)

    ax1 = fig.add_subplot(gs[0, 0])
    counts = cell_df.groupby(["assigned_cell_type", "sample"]).size().unstack(fill_value=0).loc[order]
    counts.plot(kind="bar", stacked=True, ax=ax1, color=["#4e79a7", "#f28e2b"])
    ax1.set_ylabel("Cells")
    ax1.set_xlabel("")
    ax1.set_title("A. GSE200972 LUSC tumor cells assigned by marker scores", loc="left", fontweight="bold")
    ax1.tick_params(axis="x", rotation=45)

    ax2 = fig.add_subplot(gs[0, 1])
    dot = summary.copy()
    dot["type_order"] = pd.Categorical(dot["assigned_cell_type"], categories=order[::-1], ordered=True)
    ymap = {ct: i for i, ct in enumerate(order[::-1])}
    xmap = {s: i for i, s in enumerate(dot["sample_label"].unique())}
    sc = ax2.scatter(
        dot["sample_label"].map(xmap),
        dot["assigned_cell_type"].map(ymap),
        s=np.clip(dot["PARK7_pct_detected"], 1, 100) * 7,
        c=dot["PARK7_mean_log2CPM"],
        cmap="viridis",
        edgecolor="black",
        linewidth=0.25,
    )
    ax2.set_xticks(list(xmap.values()), list(xmap.keys()), rotation=20, ha="right")
    ax2.set_yticks(list(ymap.values()), list(ymap.keys()))
    ax2.set_title("B. PARK7 mean expression and detection rate", loc="left", fontweight="bold")
    cbar = fig.colorbar(sc, ax=ax2)
    cbar.set_label("Mean PARK7 log2(CPM+1)")

    ax3 = fig.add_subplot(gs[1, 0])
    rng = np.random.default_rng(7)
    plot_df = []
    for ct in order:
        vals = cell_df.loc[cell_df["assigned_cell_type"].eq(ct), "PARK7_log2CPM"].to_numpy()
        if vals.size > 1200:
            vals = rng.choice(vals, size=1200, replace=False)
        plot_df.append(vals)
    parts = ax3.violinplot(plot_df, showmeans=False, showmedians=True, showextrema=False)
    for body, ct in zip(parts["bodies"], order):
        body.set_facecolor(colors.get(ct, "#888888"))
        body.set_alpha(0.75)
    ax3.set_xticks(np.arange(1, len(order) + 1), order, rotation=45, ha="right")
    ax3.set_ylabel("PARK7 log2(CPM+1)")
    ax3.set_title("C. PARK7 expression distribution by cell type", loc="left", fontweight="bold")

    ax4 = fig.add_subplot(gs[1, 1])
    b_summary = summary[summary["assigned_cell_type"].isin(["B cells", "Plasma cells"])].copy()
    b_summary["label"] = b_summary["sample_label"] + " / " + b_summary["assigned_cell_type"]
    ax4.barh(
        b_summary["label"],
        b_summary["PARK7_pct_detected"],
        color=[colors.get(x, "#888888") for x in b_summary["assigned_cell_type"]],
    )
    for y, row in enumerate(b_summary.itertuples()):
        ax4.text(row.PARK7_pct_detected + 1, y, f"mean={row.PARK7_mean_log2CPM:.2f}", va="center", fontsize=9)
    ax4.set_xlim(0, max(100, b_summary["PARK7_pct_detected"].max() + 15))
    ax4.set_xlabel("PARK7+ cells (%)")
    ax4.set_title("D. PARK7 detection in LUSC B-lineage cells", loc="left", fontweight="bold")

    for ax in [ax1, ax2, ax3, ax4]:
        ax.spines[["top", "right"]].set_visible(False)
    for ext in ["png", "pdf", "tif"]:
        fig.savefig(PLOT / f"Supplementary_Figure_S7_GSE200972_LUSC_scRNA_PARK7.{ext}", dpi=600, bbox_inches="tight")
    plt.close(fig)


TCGA_GENE_IDS = {
    "PARK7": "ENSG00000116288",
    "CD79A": "ENSG00000105369",
    "CD79B": "ENSG00000007312",
    "MS4A1": "ENSG00000156738",
    "CD19": "ENSG00000177455",
    "CD40": "ENSG00000101017",
    "MZB1": "ENSG00000170476",
    "JCHAIN": "ENSG00000132465",
    "CD3D": "ENSG00000167286",
    "NKG7": "ENSG00000105374",
    "LST1": "ENSG00000204482",
    "EPCAM": "ENSG00000119888",
    "KRT5": "ENSG00000186081",
    "TP63": "ENSG00000073282",
    "NFE2L2": "ENSG00000116044",
    "HMOX1": "ENSG00000100292",
    "NQO1": "ENSG00000181019",
    "SOD1": "ENSG00000142168",
    "CAT": "ENSG00000121691",
    "GPX1": "ENSG00000233276",
    "TXN": "ENSG00000136810",
    "GSR": "ENSG00000104687",
}


def run_tcga_lusc_correlation() -> pd.DataFrame:
    tpm_path = ROOT / "My_MR_Project" / "TCGA" / "TCGA-LUSC.star_tpm.tsv"
    wanted = set(TCGA_GENE_IDS.values())
    rows = []
    for chunk in pd.read_csv(tpm_path, sep="\t", chunksize=2000):
        base = chunk["Ensembl_ID"].astype(str).str.replace(r"\..*$", "", regex=True)
        hit = chunk[base.isin(wanted)].copy()
        if not hit.empty:
            hit["Ensembl_base"] = base[base.isin(wanted)].to_numpy()
            rows.append(hit)
    expr = pd.concat(rows, ignore_index=True)
    id_to_symbol = {v: k for k, v in TCGA_GENE_IDS.items()}
    expr["gene"] = expr["Ensembl_base"].map(id_to_symbol)
    expr = expr.drop(columns=["Ensembl_ID", "Ensembl_base"]).set_index("gene")
    expr = expr[~expr.index.duplicated(keep="first")]

    tumor_cols = [c for c in expr.columns if re.search(r"-01[A-Z]$", c)]
    expr = expr[tumor_cols].T
    expr.index = expr.index.str.slice(0, 12)
    expr = expr.groupby(expr.index).first()

    modules = pd.DataFrame(index=expr.index)
    b_genes = [g for g in ["CD79A", "CD79B", "MS4A1", "CD19", "CD40", "MZB1", "JCHAIN"] if g in expr]
    redox_genes = [g for g in ["NFE2L2", "HMOX1", "NQO1", "SOD1", "CAT", "GPX1", "TXN", "GSR"] if g in expr]
    squamous_genes = [g for g in ["EPCAM", "KRT5", "TP63"] if g in expr]
    modules["B_lineage_marker_module"] = expr[b_genes].mean(axis=1)
    modules["NRF2_redox_marker_module"] = expr[redox_genes].mean(axis=1)
    modules["Squamous_epithelial_module"] = expr[squamous_genes].mean(axis=1)
    modules["PARK7"] = expr["PARK7"]

    targets = [c for c in expr.columns if c != "PARK7"] + [
        "B_lineage_marker_module", "NRF2_redox_marker_module", "Squamous_epithelial_module"
    ]
    combined = pd.concat([expr, modules.drop(columns=["PARK7"])], axis=1)
    res = []
    for target in targets:
        sub = combined[["PARK7", target]].dropna()
        if len(sub) < 10:
            continue
        sp = stats.spearmanr(sub["PARK7"], sub[target])
        pr = stats.pearsonr(sub["PARK7"], sub[target])
        res.append({
            "cohort": "TCGA-LUSC",
            "n_tumor_patients": len(sub),
            "target": target,
            "target_class": (
                "module" if target.endswith("_module") else
                "B-lineage marker" if target in b_genes else
                "redox marker" if target in redox_genes else
                "other marker"
            ),
            "spearman_rho": sp.statistic,
            "spearman_p": sp.pvalue,
            "pearson_r": pr.statistic,
            "pearson_p": pr.pvalue,
        })
    out = pd.DataFrame(res)
    out["spearman_fdr"] = bh_fdr(out["spearman_p"])
    out = out.sort_values("spearman_rho", ascending=False)
    out.to_csv(TABLE / "TCGA_LUSC_PARK7_bulk_correlation.csv", index=False)
    combined.reset_index(names="TCGA_patient").to_csv(TABLE / "TCGA_LUSC_selected_expression_modules.csv", index=False)
    make_tcga_fig(out, combined)
    return out


def make_tcga_fig(out: pd.DataFrame, combined: pd.DataFrame) -> None:
    modules = ["B_lineage_marker_module", "NRF2_redox_marker_module", "Squamous_epithelial_module"]
    fig, axes = plt.subplots(1, 2, figsize=(12, 5), constrained_layout=True)
    selected = modules + ["CD79A", "MS4A1", "CD19", "CD40", "GPX1", "SOD1", "CAT", "NFE2L2", "HMOX1", "NQO1"]
    top = out[out["target"].isin(selected)].copy()
    top = top.sort_values("spearman_rho")
    colors = top["target_class"].map({
        "module": "#2f6f9f",
        "B-lineage marker": "#c74732",
        "redox marker": "#4d8f3a",
        "other marker": "#6b6b6b",
    }).fillna("#6b6b6b")
    axes[0].barh(top["target"], top["spearman_rho"], color=colors)
    axes[0].axvline(0, color="black", lw=0.8)
    axes[0].set_xlabel("Spearman rho with PARK7")
    axes[0].set_title("A. TCGA-LUSC PARK7 bulk correlations", loc="left", fontweight="bold")

    x = combined["PARK7"]
    y = combined["B_lineage_marker_module"]
    axes[1].scatter(x, y, s=18, alpha=0.65, color="#c74732", edgecolor="none")
    slope, intercept, *_ = stats.linregress(x, y)
    xs = np.linspace(x.min(), x.max(), 100)
    axes[1].plot(xs, intercept + slope * xs, color="black", lw=1)
    rho = out.loc[out["target"].eq("B_lineage_marker_module"), "spearman_rho"].iloc[0]
    p = out.loc[out["target"].eq("B_lineage_marker_module"), "spearman_p"].iloc[0]
    axes[1].set_xlabel("PARK7 TPM")
    axes[1].set_ylabel("B-lineage marker module TPM")
    axes[1].set_title(f"B. PARK7 vs B-lineage module\nrho={rho:.2f}, p={p:.2e}", loc="left", fontweight="bold")
    for ax in axes:
        ax.spines[["top", "right"]].set_visible(False)
    for ext in ["png", "pdf", "tif"]:
        fig.savefig(PLOT / f"TCGA_LUSC_PARK7_bulk_correlation.{ext}", dpi=600, bbox_inches="tight")
    plt.close(fig)


def rgb_to_dab_od(arr: np.ndarray) -> tuple[np.ndarray, np.ndarray]:
    rgb = np.clip(arr[..., :3].astype(float), 1, 255) / 255.0
    od = -np.log(rgb)
    # Approximate H-DAB deconvolution vectors. Positive DAB is brown.
    h = np.array([0.650, 0.704, 0.286])
    dab = np.array([0.268, 0.570, 0.776])
    h = h / np.linalg.norm(h)
    dab = dab / np.linalg.norm(dab)
    h_od = od @ h
    dab_od = od @ dab
    return h_od, dab_od


def quantify_ihc_image(path: Path, specimen: str, field: str) -> tuple[pd.DataFrame, dict]:
    img = Image.open(path).convert("RGB")
    max_side = 2600
    scale = min(1.0, max_side / max(img.size))
    if scale < 1.0:
        img = img.resize((int(img.size[0] * scale), int(img.size[1] * scale)))
    arr = np.asarray(img)
    h_od, dab_od = rgb_to_dab_od(arr)
    hsv_v = arr.mean(axis=2) / 255.0
    tissue = (h_od + dab_od > 0.16) & (hsv_v < 0.93)

    tile = 256
    roi_rows = []
    for y0 in range(0, arr.shape[0] - tile + 1, tile):
        for x0 in range(0, arr.shape[1] - tile + 1, tile):
            tm = tissue[y0:y0 + tile, x0:x0 + tile]
            frac = float(tm.mean())
            if frac < 0.20:
                continue
            d = dab_od[y0:y0 + tile, x0:x0 + tile][tm]
            weak = np.mean((d >= 0.08) & (d < 0.18)) * 100
            moderate = np.mean((d >= 0.18) & (d < 0.35)) * 100
            strong = np.mean(d >= 0.35) * 100
            hscore = weak + 2 * moderate + 3 * strong
            roi_rows.append({
                "image_file": path.name,
                "specimen_compartment": specimen,
                "field_roi_type": field,
                "roi_x": x0,
                "roi_y": y0,
                "tissue_fraction": frac,
                "dab_od_mean": float(np.mean(d)),
                "weak_1plus_percent": weak,
                "moderate_2plus_percent": moderate,
                "strong_3plus_percent": strong,
                "image_roi_hscore": hscore,
            })
    roi = pd.DataFrame(roi_rows)
    tissue_d = dab_od[tissue]
    weak = np.mean((tissue_d >= 0.08) & (tissue_d < 0.18)) * 100
    moderate = np.mean((tissue_d >= 0.18) & (tissue_d < 0.35)) * 100
    strong = np.mean(tissue_d >= 0.35) * 100
    summary = {
        "image_file": path.name,
        "specimen_compartment": specimen,
        "field_roi_type": field,
        "n_roi_tiles": int(len(roi)),
        "tissue_pixels": int(tissue.sum()),
        "whole_roi_dab_od_mean": float(np.mean(tissue_d)),
        "whole_roi_weak_1plus_percent": weak,
        "whole_roi_moderate_2plus_percent": moderate,
        "whole_roi_strong_3plus_percent": strong,
        "whole_roi_image_hscore": weak + 2 * moderate + 3 * strong,
        "tile_hscore_median": float(roi["image_roi_hscore"].median()) if len(roi) else np.nan,
        "tile_hscore_iqr": float(roi["image_roi_hscore"].quantile(0.75) - roi["image_roi_hscore"].quantile(0.25)) if len(roi) else np.nan,
    }
    return roi, summary


def run_ihc_roi_hscore() -> pd.DataFrame:
    img_dir = FRONTIERS_ROOT / "06_Original_Microscopy_Images"
    images = [
        ("S6C_PARK7_IHC_Biopsy_overview_original_deidentified.png", "Biopsy", "overview"),
        ("S6D_PARK7_IHC_Biopsy_matched_field_original_deidentified.png", "Biopsy", "matched_field"),
        ("S6G_PARK7_IHC_Surgery_overview_original_deidentified.png", "Surgery", "overview"),
        ("S6H_PARK7_IHC_Surgery_matched_field_original_deidentified.png", "Surgery", "matched_field"),
    ]
    roi_frames = []
    summaries = []
    for name, specimen, field in images:
        roi, summary = quantify_ihc_image(img_dir / name, specimen, field)
        roi_frames.append(roi)
        summaries.append(summary)
    roi_all = pd.concat(roi_frames, ignore_index=True)
    summary_df = pd.DataFrame(summaries)
    roi_all.to_csv(TABLE / "IHC_PARK7_ROI_based_image_Hscore_tiles.csv", index=False)
    summary_df.to_csv(TABLE / "IHC_PARK7_compartment_aware_ROI_Hscore_summary.csv", index=False)

    reader_path = FRONTIERS_ROOT / "05_Source_Data" / "Source_Data_CSV" / "SupplementaryTableS4_IHC_Hscore_Details_Deidentified_Source.csv"
    reader = pd.read_csv(reader_path)
    reader_summary = reader.groupby("Specimen_type", as_index=False).agg(
        n_cases=("Case_ID", "size"),
        mean_reader_hscore=("Mean_H_score", "mean"),
        sd_reader_hscore=("Mean_H_score", "std"),
        median_reader_hscore=("Mean_H_score", "median"),
    )
    if reader["Specimen_type"].nunique() == 2:
        groups = [g["Mean_H_score"].to_numpy() for _, g in reader.groupby("Specimen_type")]
        mw = stats.mannwhitneyu(groups[0], groups[1], alternative="two-sided")
        reader_summary["biopsy_vs_surgery_mannwhitney_p"] = mw.pvalue
    reader_summary.to_csv(TABLE / "IHC_PARK7_reader_Hscore_by_specimen_compartment.csv", index=False)
    make_ihc_fig(roi_all, summary_df, reader)
    return summary_df


def make_ihc_fig(roi_all: pd.DataFrame, summary_df: pd.DataFrame, reader: pd.DataFrame) -> None:
    fig, axes = plt.subplots(1, 3, figsize=(15.5, 4.8), constrained_layout=True)
    labels = summary_df["specimen_compartment"] + " " + summary_df["field_roi_type"]
    axes[0].bar(labels, summary_df["whole_roi_image_hscore"], color=["#c74732", "#ef8a47", "#2f6f9f", "#4e79a7"])
    axes[0].tick_params(axis="x", rotation=35)
    axes[0].set_ylabel("Image-derived H-score")
    axes[0].set_title("A. ROI-level PARK7 IHC quantification", loc="left", fontweight="bold")

    data = [
        roi_all.loc[roi_all["specimen_compartment"].eq("Biopsy"), "image_roi_hscore"],
        roi_all.loc[roi_all["specimen_compartment"].eq("Surgery"), "image_roi_hscore"],
    ]
    axes[1].boxplot(data, labels=["Biopsy", "Surgery"], patch_artist=True)
    axes[1].set_ylabel("Tile H-score")
    axes[1].set_title("B. ROI-tile distribution by compartment", loc="left", fontweight="bold")

    reader_groups = [
        reader.loc[reader["Specimen_type"].eq("Biopsy"), "Mean_H_score"].dropna(),
        reader.loc[reader["Specimen_type"].eq("Surgery"), "Mean_H_score"].dropna(),
    ]
    axes[2].boxplot(reader_groups, labels=["Biopsy", "Surgery"], patch_artist=True)
    axes[2].set_title("C. Reader H-score", loc="left", fontweight="bold")
    axes[2].set_xlabel("")
    axes[2].set_ylabel("Mean reader H-score")
    fig.suptitle("")
    for ax in axes:
        ax.spines[["top", "right"]].set_visible(False)
    for ext in ["png", "pdf", "tif"]:
        fig.savefig(PLOT / f"IHC_PARK7_compartment_ROI_Hscore.{ext}", dpi=600, bbox_inches="tight")
    plt.close(fig)


def update_submission_packages() -> None:
    mappings = [
        (PLOT / "Supplementary_Figure_S7_GSE200972_LUSC_scRNA_PARK7.tif",
         FRONTIERS_ROOT / "04_Supplementary_Material" / "S7.tif"),
        (PLOT / "Supplementary_Figure_S7_GSE200972_LUSC_scRNA_PARK7.tif",
         FRONTIERS_68 / "Supplementary_File_PARK7_LUSC" / "S7.tif"),
        (PLOT / "TCGA_LUSC_PARK7_bulk_correlation.tif",
         FRONTIERS_ROOT / "04_Supplementary_Material" / "S8_TCGA_LUSC_bulk_correlation.tif"),
        (PLOT / "IHC_PARK7_compartment_ROI_Hscore.tif",
         FRONTIERS_ROOT / "04_Supplementary_Material" / "S9_IHC_ROI_Hscore.tif"),
    ]
    for src, dst in mappings:
        if src.exists() and dst.parent.exists():
            shutil.copy2(src, dst)

    source_dirs = [
        FRONTIERS_ROOT / "05_Source_Data" / "Source_Data_CSV",
        FRONTIERS_68 / "Source_Data_CSV",
    ]
    table_files = [
        "S7_GSE200972_LUSC_scRNA_PARK7_celltype_summary.csv",
        "S7_GSE200972_LUSC_scRNA_PARK7_cell_level.csv",
        "SupplementaryTableS5_CMap_Top50_Filtered_Signature_Source.csv",
        "CMap_Input_Up_top50_filtered.txt",
        "CMap_Input_Down_top50_filtered.txt",
        "CMap_Input_Tags_top50_filtered.txt",
        "TCGA_LUSC_PARK7_bulk_correlation.csv",
        "TCGA_LUSC_selected_expression_modules.csv",
        "IHC_PARK7_ROI_based_image_Hscore_tiles.csv",
        "IHC_PARK7_compartment_aware_ROI_Hscore_summary.csv",
        "IHC_PARK7_reader_Hscore_by_specimen_compartment.csv",
    ]
    for sdir in source_dirs:
        if not sdir.exists():
            continue
        for name in table_files:
            src = TABLE / name
            if src.exists():
                shutil.copy2(src, sdir / name)

    update_manifest_files()
    update_supplementary_docx_notes()

    rebuild_zip(FRONTIERS_ROOT / "04_Supplementary_Material.zip", FRONTIERS_ROOT / "04_Supplementary_Material")
    rebuild_zip(FRONTIERS_ROOT / "05_Source_Data" / "Source_Data_CSV.zip", FRONTIERS_ROOT / "05_Source_Data" / "Source_Data_CSV")
    rebuild_zip(FRONTIERS_ROOT / "05_Source_Data.zip", FRONTIERS_ROOT / "05_Source_Data")
    rebuild_zip(FRONTIERS_ROOT.with_suffix(".zip"), FRONTIERS_ROOT)
    if FRONTIERS_68.exists():
        rebuild_zip(FRONTIERS_68 / "Supplementary_File_PARK7_LUSC.zip", FRONTIERS_68 / "Supplementary_File_PARK7_LUSC")
        rebuild_zip(FRONTIERS_68 / "Source_Data_CSV.zip", FRONTIERS_68 / "Source_Data_CSV")


def update_manifest_files() -> None:
    new_source_rows = [
        ("S7_GSE200972_LUSC_scRNA_PARK7_celltype_summary.csv", "Supplementary Fig. S7", "Public LUSC scRNA-seq PARK7 expression summary by marker-score-assigned cell type."),
        ("S7_GSE200972_LUSC_scRNA_PARK7_cell_level.csv", "Supplementary Fig. S7", "Cell-level PARK7 log2(CPM+1), detection status, sample and assigned cell type for GSE200972 LUSC tumor samples."),
        ("SupplementaryTableS5_CMap_Top50_Filtered_Signature_Source.csv", "Supplementary Table S5 / CMap input", "Top 50 up and top 50 down genes from LUSC B-lineage PARK7-high vs PARK7-low cells after excluding PARK7/NRF2/redox anchor genes."),
        ("CMap_Input_Up_top50_filtered.txt", "CMap input", "Filtered top 50 upregulated genes for CLUE/CMap upload."),
        ("CMap_Input_Down_top50_filtered.txt", "CMap input", "Filtered top 50 downregulated genes for CLUE/CMap upload."),
        ("CMap_Input_Tags_top50_filtered.txt", "CMap input", "Combined up/down filtered CMap tag file."),
        ("TCGA_LUSC_PARK7_bulk_correlation.csv", "Supplementary Fig. S8", "TCGA-LUSC PARK7 bulk-expression correlations with B-lineage, redox and squamous marker/module variables."),
        ("TCGA_LUSC_selected_expression_modules.csv", "Supplementary Fig. S8", "Selected TCGA-LUSC expression and derived module source data used for PARK7 correlation plots."),
        ("IHC_PARK7_ROI_based_image_Hscore_tiles.csv", "Supplementary Fig. S9", "Tile-level image-derived PARK7 DAB optical-density H-score values from biopsy/surgery overview and matched-field ROIs."),
        ("IHC_PARK7_compartment_aware_ROI_Hscore_summary.csv", "Supplementary Fig. S9", "Compartment-aware ROI-level PARK7 image H-score summary for representative biopsy and surgery IHC images."),
        ("IHC_PARK7_reader_Hscore_by_specimen_compartment.csv", "Supplementary Fig. S9 / Table S4", "Reader/pathologist H-score summary by specimen compartment with biopsy-vs-surgery nonparametric comparison."),
    ]
    for manifest in [
        FRONTIERS_ROOT / "05_Source_Data" / "Source_Data_CSV" / "Source_Data_CSV_Manifest.csv",
        FRONTIERS_68 / "Source_Data_CSV" / "Source_Data_CSV_Manifest.csv",
    ]:
        if not manifest.exists():
            continue
        df = pd.read_csv(manifest)
        existing = set(df["file_name"].astype(str))
        add = [r for r in new_source_rows if r[0] not in existing]
        if add:
            df = pd.concat([df, pd.DataFrame(add, columns=df.columns[:3])], ignore_index=True)
            df.to_csv(manifest, index=False)

    package_manifest = FRONTIERS_ROOT / "PACKAGE_MANIFEST.csv"
    if package_manifest.exists():
        rows = []
        for path in FRONTIERS_ROOT.rglob("*"):
            if not path.is_file():
                continue
            if path.name.endswith(".bak_before_final_analyses"):
                continue
            rel = path.relative_to(FRONTIERS_ROOT).as_posix()
            rows.append({
                "relative_path": rel,
                "size_bytes": path.stat().st_size,
                "upload_role": package_role(rel),
            })
        pd.DataFrame(rows).sort_values("relative_path").to_csv(package_manifest, index=False)

    readme = FRONTIERS_ROOT / "README_Frontiers_Submission_Checklist.md"
    if readme.exists():
        text = readme.read_text(encoding="utf-8")
        marker = "## Final Analysis Addendum"
        if marker not in text:
            text += (
                "\n\n## Final Analysis Addendum\n\n"
                "- Added Supplementary Fig. S7: public LUSC scRNA-seq PARK7 cell-type validation from GSE200972.\n"
                "- Added Supplementary Fig. S8: TCGA-LUSC PARK7 bulk-expression correlation analysis.\n"
                "- Added Supplementary Fig. S9: compartment-aware / ROI-based PARK7 IHC H-score quantification.\n"
                "- Added filtered CMap top 50 up/down signature files excluding PARK7/NRF2/redox anchor genes.\n"
            )
            readme.write_text(text, encoding="utf-8")


def package_role(rel: str) -> str:
    if rel.startswith("01_Manuscript/"):
        return "core manuscript"
    if rel.startswith("02_Frontiers_Statements/"):
        return "Frontiers submission textbox support"
    if rel.startswith("03_Figures_Main") or rel.startswith("03_Figures_Main/"):
        return "main figure upload"
    if rel.startswith("04_Supplementary_Material") or rel.startswith("04_Supplementary_Material/"):
        return "supplementary material"
    if rel.startswith("05_Source_Data") or rel.startswith("05_Source_Data/"):
        return "source data / reproducibility"
    if rel.startswith("06_Original_Microscopy_Images") or rel.startswith("06_Original_Microscopy_Images/"):
        return "supporting original microscopy images"
    if rel.startswith("07_Admin_and_Optional/"):
        return "admin optional"
    return "package documentation"


def update_supplementary_docx_notes() -> None:
    try:
        from docx import Document
    except Exception:
        return
    docx_paths = [
        FRONTIERS_ROOT / "04_Supplementary_Material" / "Supplementary_Material.docx",
        FRONTIERS_68 / "Supplementary_File_PARK7_LUSC" / "Supplementary_Material.docx",
    ]
    additions = [
        ("Supplementary Figure S7. Public LUSC scRNA-seq validation of PARK7 expression.",
         "GSE200972 tumor samples GSM6047623_P1_T_R_I and GSM6047625_P1_T_R_M were analyzed from the public GEO raw matrices. Major cell classes were assigned using canonical marker-score annotation, and PARK7 expression was summarized as log2(CPM+1) and detection percentage across cell types, with specific B-lineage summaries shown for B cells and plasma cells."),
        ("Supplementary Figure S8. TCGA-LUSC bulk correlation analysis.",
         "TCGA-LUSC tumor RNA-seq TPM profiles were used to correlate PARK7 expression with selected B-lineage, NRF2/redox and squamous epithelial marker/module variables. Correlations are reported with Spearman and Pearson statistics and FDR-adjusted P values in the source-data CSV files."),
        ("Supplementary Figure S9. PARK7 IHC compartment-aware ROI H-score quantification.",
         "Representative biopsy and surgery PARK7 IHC overview and matched-field images were quantified by tile/ROI-level DAB optical-density H-score, and reader/pathologist H-score distributions were summarized by specimen compartment. These data support tissue-level PARK7 staining assessment and are not interpreted as B-cell-specific protein localization."),
        ("Supplementary Table S5A. Filtered CMap top 50 up/down signature.",
         "The CMap query signature was regenerated from public LUSC B-lineage scRNA-seq cells stratified by PARK7 expression. The top 50 upregulated and top 50 downregulated genes were selected after excluding PARK7/NRF2/redox anchor genes and low-detection genes, to reduce circularity from manually specified pathway genes."),
    ]
    for path in docx_paths:
        if not path.exists():
            continue
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        if "Supplementary Figure S7. Public LUSC scRNA-seq validation" in text:
            continue
        backup = path.with_suffix(path.suffix + ".bak_before_final_analyses")
        if not backup.exists():
            shutil.copy2(path, backup)
        doc.add_paragraph("")
        for title, body in additions:
            p = doc.add_paragraph()
            p.add_run(title).bold = True
            doc.add_paragraph(body)
        doc.save(path)


def rebuild_zip(zip_path: Path, source_dir: Path) -> None:
    if not source_dir.exists():
        return
    if zip_path.exists():
        backup_root = WORK / "zip_backups_removed_from_upload_folder"
        backup_root.mkdir(parents=True, exist_ok=True)
        backup = backup_root / (zip_path.name + ".bak_before_final_analyses")
        if not backup.exists():
            shutil.copy2(zip_path, backup)
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for path in source_dir.rglob("*"):
            if path.is_file():
                if path.name.endswith(".bak_before_final_analyses"):
                    continue
                z.write(path, path.relative_to(source_dir.parent))


def write_summary_report(sc_summary, cmap_sig, tcga_corr, ihc_summary) -> None:
    report = []
    report.append("# Final unfinished analyses completion report")
    report.append("")
    report.append("Completed items:")
    report.append("- Public LUSC scRNA-seq validation using GSE200972 tumor samples GSM6047623 and GSM6047625.")
    report.append("- TCGA-LUSC bulk correlation of PARK7 with B-lineage, redox, and epithelial markers/modules.")
    report.append("- IHC compartment-aware reader H-score summary and ROI/tile-based image quantification.")
    report.append("- CMap top 50 up/down signature from LUSC B-lineage PARK7-high vs PARK7-low cells, excluding PARK7/NRF2/redox anchor genes.")
    report.append("- Real Supplementary Figure S7 scRNA panel exported as PNG/PDF/TIF.")
    report.append("")
    report.append(f"GSE200972 assigned cell-type rows: {len(sc_summary)} summary rows.")
    report.append(f"CMap signature genes: {cmap_sig['CMap_direction'].value_counts().to_dict()}.")
    report.append(f"TCGA-LUSC correlations tested: {len(tcga_corr)}.")
    report.append(f"IHC ROI images quantified: {len(ihc_summary)}.")
    report.append("")
    report.append("Key output directory:")
    report.append(str(OUT))
    (WORK / "FINAL_ANALYSES_COMPLETION_REPORT.md").write_text("\n".join(report), encoding="utf-8")


def main() -> None:
    warnings.filterwarnings("ignore")
    sc_cell, sc_summary, cmap_sig = run_public_lusc_scrna()
    tcga_corr = run_tcga_lusc_correlation()
    ihc_summary = run_ihc_roi_hscore()
    update_submission_packages()
    write_summary_report(sc_summary, cmap_sig, tcga_corr, ihc_summary)
    print("Completed final unfinished analyses.")
    print(WORK / "FINAL_ANALYSES_COMPLETION_REPORT.md")


if __name__ == "__main__":
    main()
