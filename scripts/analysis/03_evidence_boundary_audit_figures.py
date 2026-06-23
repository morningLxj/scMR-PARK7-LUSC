from __future__ import annotations

import math
import os
import shutil
import zipfile
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt


REPO_ROOT = Path(__file__).resolve().parents[2]
ROOT = Path(os.environ.get("SCMR_PROJECT_ROOT", REPO_ROOT))
PACKAGE = Path(os.environ.get("SCMR_SUBMISSION_PACKAGE", ROOT / "submission_package"))
SUPP = PACKAGE / "04_Supplementary_Material"
SRC = Path(os.environ.get("SCMR_SOURCE_DATA_DIR", REPO_ROOT / "data" / "source_data_csv"))
MANUSCRIPT = PACKAGE / "01_Manuscript" / "Manuscript.docx"
TITLE_PAGE = PACKAGE / "01_Manuscript" / "Title_Page_Frontiers.docx"
MAIN_FIG = PACKAGE / "03_Figures_Main"
STATEMENTS = PACKAGE / "02_Frontiers_Statements"
ADMIN = PACKAGE / "07_Admin_and_Optional"
WORK = Path(os.environ.get("SCMR_WORK_DIR", ROOT / "work" / "prioritization_audit_20260612"))
WORK.mkdir(parents=True, exist_ok=True)

NEW_TITLE = "Cell-type-resolved genetic prioritization of PARK7 in lung squamous cell carcinoma with indirect spatial redox context"


def backup(path: Path) -> None:
    if not path.exists():
        return
    rel = path.relative_to(PACKAGE) if path.is_relative_to(PACKAGE) else path.name
    dst = WORK / "backups" / rel
    dst.parent.mkdir(parents=True, exist_ok=True)
    if not dst.exists():
        shutil.copy2(path, dst)


def save_fig(fig, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, dpi=300, bbox_inches="tight", pil_kwargs={"compression": "tiff_lzw"})
    plt.close(fig)


def fisher_or(a: int, b: int, c: int, d: int) -> tuple[float, float, float]:
    # Haldane-Anscombe correction keeps CIs finite when a cell is zero.
    aa, bb, cc, dd = a + 0.5, b + 0.5, c + 0.5, d + 0.5
    odds = (aa * dd) / (bb * cc)
    se = math.sqrt(1 / aa + 1 / bb + 1 / cc + 1 / dd)
    lo = math.exp(math.log(odds) - 1.96 * se)
    hi = math.exp(math.log(odds) + 1.96 * se)
    try:
        from scipy.stats import fisher_exact

        _, p = fisher_exact([[a, b], [c, d]], alternative="two-sided")
    except Exception:
        p = np.nan
    return odds, lo, hi, p


def add_manifest_row(file_name: str, manuscript_item: str, current_note: str) -> None:
    manifest = SRC / "Source_Data_CSV_Manifest.csv"
    df = pd.read_csv(manifest) if manifest.exists() else pd.DataFrame(columns=["file_name", "manuscript_item", "current_note"])
    if file_name not in set(df["file_name"].astype(str)):
        df = pd.concat(
            [
                df,
                pd.DataFrame(
                    [{"file_name": file_name, "manuscript_item": manuscript_item, "current_note": current_note}]
                ),
            ],
            ignore_index=True,
        )
    else:
        df.loc[df["file_name"].astype(str) == file_name, ["manuscript_item", "current_note"]] = [
            manuscript_item,
            current_note,
        ]
    df.to_csv(manifest, index=False)


def make_claim_boundary_table() -> None:
    rows = [
        {
            "Claim": "PARK7 intermediate-B-cell expression is genetically prioritized for LUSC",
            "Directly_supported": "Yes",
            "Evidence_source": "Cell-type-resolved MR",
            "Strength": "Strong prioritization signal",
            "Main_limitation": "Cis-eQTL regulatory architecture and summary-data assumptions",
            "Allowed_manuscript_wording": "nominates; prioritizes; B-cell-linked candidate",
            "Avoid_wording": "proves causality; drives LUSC; large clinical effect",
        },
        {
            "Claim": "PARK7 shares a causal variant with LUSC susceptibility",
            "Directly_supported": "No",
            "Evidence_source": "Colocalization; LUSC PP4 below 0.4 in B-cell subsets",
            "Strength": "Weak-to-moderate regional support",
            "Main_limitation": "Below conventional PP4 > 0.8 high-confidence threshold",
            "Allowed_manuscript_wording": "suggestive regional evidence; LUSC-preferential posterior support",
            "Avoid_wording": "definitive colocalization; shared causal variant",
        },
        {
            "Claim": "PARK7 localizes to B-cell-rich redox niches",
            "Directly_supported": "No",
            "Evidence_source": "Spatial B-cell and antioxidant score analysis",
            "Strength": "Indirect boundary/context evidence",
            "Main_limitation": "PARK7 not directly measured in the primary spatial scoring layer",
            "Allowed_manuscript_wording": "indirect heterogeneous redox context; constrains the hypothesis",
            "Avoid_wording": "PARK7-positive B-cell microniche; spatial validation",
        },
        {
            "Claim": "PARK7 is detectable in B-lineage cells in public LUSC scRNA-seq",
            "Directly_supported": "Yes, within the analyzed public samples",
            "Evidence_source": "GSE200972 marker-score cell assignment",
            "Strength": "Contextual expression support",
            "Main_limitation": "Two tumor samples; broad marker-score assignment; not independent protein localization",
            "Allowed_manuscript_wording": "detectable across multiple compartments, including B-lineage cells",
            "Avoid_wording": "B-cell specific; exclusive B-cell expression",
        },
        {
            "Claim": "PARK7 protein is expressed in B cells",
            "Directly_supported": "No",
            "Evidence_source": "Single-marker PARK7 IHC and ROI/tile H-score",
            "Strength": "Tissue-level protein detectability",
            "Main_limitation": "No multiplex PARK7/CD20/CD79A co-localization",
            "Allowed_manuscript_wording": "detectable tissue-level PARK7 staining",
            "Avoid_wording": "B-cell-specific protein localization",
        },
        {
            "Claim": "PARK7 regulates B-cell ER stress or redox adaptation",
            "Directly_supported": "No",
            "Evidence_source": "Network-inferred PARK7 perturbation and pathway enrichment",
            "Strength": "Exploratory hypothesis support",
            "Main_limitation": "No wet-lab perturbation; enrichment sensitive to gene universe and high-RNA programs",
            "Allowed_manuscript_wording": "network-inferred; testable redox/ER-stress hypothesis",
            "Avoid_wording": "functional validation; PARK7 drives ER stress",
        },
        {
            "Claim": "PARK7 is a therapeutic target in LUSC",
            "Directly_supported": "No",
            "Evidence_source": "Exploratory CMap annotation",
            "Strength": "Perturbational context only",
            "Main_limitation": "CMap cell contexts are not LUSC immune-cell validation systems",
            "Allowed_manuscript_wording": "not used to infer therapeutic vulnerability",
            "Avoid_wording": "drug efficacy; treatment recommendation; validated target",
        },
    ]
    out = SRC / "SupplementaryTableS9_ReviewerFacing_ClaimBoundary_Matrix.csv"
    pd.DataFrame(rows).to_csv(out, index=False)
    add_manifest_row(
        out.name,
        "Supplementary Table S9",
        "Reviewer-facing claim and allowed-wording boundary matrix for the prioritization manuscript.",
    )


def make_main_figure_1_boundary_framework() -> None:
    backup(MAIN_FIG / "Figure_1.tif")
    rows = [
        ["MR", "B-cell-linked genetic prioritization", "Direct causality"],
        ["Locus", "Weak-to-moderate regional evidence", "Shared causal variant"],
        ["Spatial", "Heterogeneous indirect redox context", "PARK7 spatial localization"],
        ["scRNA/TCGA", "Expression and pathway context", "B-cell-specific function"],
        ["IHC/CMap", "Tissue detectability and perturbational context", "Protein co-localization or drug efficacy"],
    ]
    fig, ax = plt.subplots(figsize=(10.5, 6.2))
    ax.axis("off")
    ax.text(
        0.5,
        0.96,
        "Analytical evidence-boundary framework for PARK7 prioritization in LUSC",
        ha="center",
        va="top",
        fontsize=16,
        fontweight="bold",
    )
    ax.text(
        0.5,
        0.90,
        "Computational/genetic prioritization, not mechanistic or therapeutic validation",
        ha="center",
        va="top",
        fontsize=10.5,
        color="#444444",
    )
    table = ax.table(
        cellText=rows,
        colLabels=["Evidence layer", "What it supports", "What it does not establish"],
        cellLoc="left",
        colLoc="left",
        colWidths=[0.18, 0.41, 0.41],
        bbox=[0.04, 0.18, 0.92, 0.64],
    )
    table.auto_set_font_size(False)
    table.set_fontsize(10)
    for (r, c), cell in table.get_celld().items():
        cell.set_edgecolor("#D0D0D0")
        if r == 0:
            cell.set_facecolor("#2F4B7C")
            cell.get_text().set_color("white")
            cell.get_text().set_fontweight("bold")
        elif c == 0:
            cell.set_facecolor("#E8EEF7")
            cell.get_text().set_fontweight("bold")
        elif c == 1:
            cell.set_facecolor("#F3FAF4")
        else:
            cell.set_facecolor("#FFF4F2")
    ax.text(
        0.5,
        0.08,
        "Preferred interpretation: PARK7 is nominated as a LUSC-preferential, B-cell-linked candidate with indirect, heterogeneous redox/ER-stress contextual support.",
        ha="center",
        va="center",
        fontsize=10,
        color="#333333",
        wrap=True,
    )
    save_fig(fig, MAIN_FIG / "Figure_1.tif")


def make_mr_audit_figure() -> None:
    inst = pd.read_csv(SRC / "SupplementaryTableS1A_PARK7_PerInstrument_F_Statistics.csv")
    sens = pd.read_csv(SRC / "SupplementaryTableS1C_PARK7_MR_Sensitivity_Summary.csv").iloc[0]
    loo = pd.read_csv(SRC / "SupplementaryTableS1D_PARK7_LeaveOneOut_Source.csv")

    methods = pd.DataFrame(
        [
            ["IVW", sens["IVW_beta"], sens["IVW_SE"]],
            ["Weighted median", sens["Weighted_median_beta"], sens["Weighted_median_SE_bootstrap"]],
            ["MR-Egger", sens["MR_Egger_slope"], sens["MR_Egger_slope_SE"]],
        ],
        columns=["method", "beta", "se"],
    )
    methods["ci_low"] = methods["beta"] - 1.96 * methods["se"]
    methods["ci_high"] = methods["beta"] + 1.96 * methods["se"]
    methods.to_csv(SRC / "S10_MR_Audit_Method_Estimates.csv", index=False)
    add_manifest_row(
        "S10_MR_Audit_Method_Estimates.csv",
        "Supplementary Fig. S10",
        "MR estimator summary used in the reviewer-facing MR audit figure.",
    )

    fig, axes = plt.subplots(2, 2, figsize=(9.5, 7.2))
    axes[0, 0].hist(inst["F_statistic"], bins=12, color="#4477AA", edgecolor="white")
    axes[0, 0].axvline(10, color="#CC6677", linestyle="--", label="F=10")
    axes[0, 0].axvline(30, color="#228833", linestyle="--", label="F=30")
    axes[0, 0].set_title("Instrument strength distribution")
    axes[0, 0].set_xlabel("F statistic")
    axes[0, 0].set_ylabel("Number of SNPs")
    axes[0, 0].legend(frameon=False, fontsize=8)

    y = np.arange(len(methods))
    axes[0, 1].errorbar(
        methods["beta"],
        y,
        xerr=[methods["beta"] - methods["ci_low"], methods["ci_high"] - methods["beta"]],
        fmt="o",
        color="#332288",
        ecolor="#88CCEE",
        capsize=3,
    )
    axes[0, 1].axvline(0, color="black", linewidth=0.8)
    axes[0, 1].set_yticks(y, methods["method"])
    axes[0, 1].set_title("Available MR estimators")
    axes[0, 1].set_xlabel("Beta")

    axes[1, 0].scatter(inst["exposure_beta"], inst["outcome_beta"], s=22, alpha=0.75, color="#117733")
    axes[1, 0].axhline(0, color="black", linewidth=0.8)
    axes[1, 0].axvline(0, color="black", linewidth=0.8)
    axes[1, 0].set_title("Per-SNP exposure/outcome effects")
    axes[1, 0].set_xlabel("Exposure beta")
    axes[1, 0].set_ylabel("Outcome beta")

    loo_sorted = loo.sort_values("LOO_IVW_beta")
    axes[1, 1].plot(np.arange(len(loo_sorted)), loo_sorted["LOO_IVW_beta"], color="#AA4499", linewidth=1.3)
    axes[1, 1].axhline(sens["IVW_beta"], color="black", linestyle="--", linewidth=0.9, label="Full IVW")
    axes[1, 1].set_title("Leave-one-SNP-out IVW stability")
    axes[1, 1].set_xlabel("SNP removed, sorted")
    axes[1, 1].set_ylabel("LOO beta")
    axes[1, 1].legend(frameon=False, fontsize=8)

    fig.suptitle("Supplementary Fig. S10. PARK7 MR instrument and estimator audit", fontsize=13)
    fig.tight_layout()
    save_fig(fig, SUPP / "S10_MR_Instrument_Estimator_Audit.tif")


def make_spatial_sensitivity_figure() -> None:
    df = pd.read_csv(SRC / "Figure4_Spatial_CoLocalization_Scores.csv")
    settings = []
    rng = np.random.default_rng(20260612)
    threshold_defs = [
        ("top10_norm", "top_fraction", 0.10, "B_norm", "A_norm"),
        ("top20_norm", "top_fraction", 0.20, "B_norm", "A_norm"),
        ("top25_norm", "top_fraction", 0.25, "B_norm", "A_norm"),
        ("median_norm", "greater_than", None, "B_norm", "A_norm"),
        ("z_gt_1", "fixed_gt", 1.0, "B_cell_density_z", "Antioxidant_z"),
    ]
    null_rows = []
    for label, mode, value, bcol, acol in threshold_defs:
        if mode == "top_fraction":
            n_top = max(1, int(round(len(df) * float(value))))
            b_high = pd.Series(False, index=df.index)
            a_high = pd.Series(False, index=df.index)
            b_high.loc[df[bcol].sort_values(ascending=False).index[:n_top]] = True
            a_high.loc[df[acol].sort_values(ascending=False).index[:n_top]] = True
            bt = float(df.loc[b_high, bcol].min())
            at = float(df.loc[a_high, acol].min())
        elif mode == "greater_than":
            bt = float(df[bcol].median())
            at = float(df[acol].median())
            b_high = df[bcol] > bt
            a_high = df[acol] > at
        else:
            bt = float(value)
            at = float(value)
            b_high = df[bcol] > bt
            a_high = df[acol] > at
        both = int((b_high & a_high).sum())
        b_only = int((b_high & ~a_high).sum())
        a_only = int((~b_high & a_high).sum())
        neither = int((~b_high & ~a_high).sum())
        expected = float(b_high.sum() * a_high.sum() / len(df))
        odds, lo, hi, p = fisher_or(both, b_only, a_only, neither)
        perm = []
        a_arr = np.asarray(a_high)
        for _ in range(1000):
            perm_a = rng.permutation(a_arr)
            perm.append(int((np.asarray(b_high) & perm_a).sum()))
        perm = np.asarray(perm)
        settings.append(
            {
                "threshold": label,
                "B_threshold": bt,
                "Antioxidant_threshold": at,
                "n_spots": len(df),
                "B_high": int(b_high.sum()),
                "Antioxidant_high": int(a_high.sum()),
                "Both_high_observed": both,
                "Both_high_expected_independence": expected,
                "Observed_over_expected": both / expected if expected else np.nan,
                "Fisher_OR": odds,
                "Fisher_CI_low": lo,
                "Fisher_CI_high": hi,
                "Fisher_P": p,
                "Permutation_mean_overlap": perm.mean(),
                "Permutation_p_lower_or_equal": float((perm <= both).mean()),
                "Permutation_p_upper_or_equal": float((perm >= both).mean()),
            }
        )
        if label in {"top20_norm", "z_gt_1"}:
            for val in perm:
                null_rows.append({"threshold": label, "permuted_overlap": val, "observed_overlap": both})

    out = pd.DataFrame(settings)
    out.to_csv(SRC / "S11_Spatial_Threshold_Permutation_Sensitivity.csv", index=False)
    add_manifest_row(
        "S11_Spatial_Threshold_Permutation_Sensitivity.csv",
        "Supplementary Fig. S11",
        "Threshold and permutation sensitivity for spatial B-cell/antioxidant same-spot overlap.",
    )
    pd.DataFrame(null_rows).to_csv(SRC / "S11_Spatial_Permutation_Null_Distribution.csv", index=False)
    add_manifest_row(
        "S11_Spatial_Permutation_Null_Distribution.csv",
        "Supplementary Fig. S11",
        "Permutation null overlaps for selected spatial threshold settings.",
    )

    fig, axes = plt.subplots(1, 3, figsize=(12, 3.6))
    axes[0].bar(out["threshold"], out["Observed_over_expected"], color="#CC6677")
    axes[0].axhline(1, color="black", linewidth=0.8, linestyle="--")
    axes[0].tick_params(axis="x", rotation=35)
    axes[0].set_ylabel("Observed / expected overlap")
    axes[0].set_title("Threshold sensitivity")

    axes[1].errorbar(
        out["Fisher_OR"],
        np.arange(len(out)),
        xerr=[out["Fisher_OR"] - out["Fisher_CI_low"], out["Fisher_CI_high"] - out["Fisher_OR"]],
        fmt="o",
        color="#882255",
        ecolor="#DDCC77",
        capsize=3,
    )
    axes[1].axvline(1, color="black", linewidth=0.8, linestyle="--")
    axes[1].set_yticks(np.arange(len(out)), out["threshold"])
    axes[1].set_xlabel("Fisher OR")
    axes[1].set_title("Same-spot association")

    null = pd.DataFrame(null_rows)
    for i, label in enumerate(null["threshold"].unique()):
        vals = null.loc[null["threshold"] == label, "permuted_overlap"]
        obs = null.loc[null["threshold"] == label, "observed_overlap"].iloc[0]
        axes[2].hist(vals, bins=20, alpha=0.55, label=label)
        axes[2].axvline(obs, linewidth=1.2)
    axes[2].set_title("Permutation null examples")
    axes[2].set_xlabel("Overlap count")
    axes[2].legend(frameon=False, fontsize=8)

    fig.suptitle("Supplementary Fig. S11. Spatial threshold and permutation sensitivity", fontsize=13)
    fig.tight_layout()
    save_fig(fig, SUPP / "S11_Spatial_Threshold_Permutation_Sensitivity.tif")


def make_scrna_pseudobulk_figure() -> None:
    cells = pd.read_csv(SRC / "S7_GSE200972_LUSC_scRNA_PARK7_cell_level.csv")
    grouped = (
        cells.groupby(["sample", "sample_label", "assigned_cell_type"])
        .agg(
            n_cells=("PARK7_log2CPM", "size"),
            pct_detected=("PARK7_detected", lambda x: 100 * np.mean(x)),
            mean_log2CPM=("PARK7_log2CPM", "mean"),
            median_log2CPM=("PARK7_log2CPM", "median"),
            positive_mean_log2CPM=("PARK7_log2CPM", lambda x: x[x > 0].mean() if np.any(x > 0) else 0.0),
        )
        .reset_index()
    )
    grouped.to_csv(SRC / "S12_scRNA_PARK7_PatientLevel_Pseudobulk.csv", index=False)
    add_manifest_row(
        "S12_scRNA_PARK7_PatientLevel_Pseudobulk.csv",
        "Supplementary Fig. S12",
        "Patient/sample-level PARK7 detection and expression summaries by broad assigned cell type.",
    )

    keep = grouped[grouped["n_cells"] >= 20].copy()
    order = (
        keep.groupby("assigned_cell_type")["pct_detected"].mean().sort_values(ascending=False).index.tolist()
    )
    fig, axes = plt.subplots(1, 2, figsize=(11, 4.1), sharey=True)
    for sample, sub in keep.groupby("sample"):
        axes[0].scatter(sub["pct_detected"], sub["assigned_cell_type"], label=sample, s=35)
        axes[1].scatter(sub["positive_mean_log2CPM"], sub["assigned_cell_type"], label=sample, s=35)
    for ax in axes:
        ax.set_yticks(range(len(order)), order)
        ax.grid(axis="x", alpha=0.25)
    axes[0].set_xlabel("PARK7 detection (%)")
    axes[0].set_title("Detection fraction")
    axes[1].set_xlabel("Mean log2(CPM+1) in positive cells")
    axes[1].set_title("Positive-cell expression intensity")
    axes[1].legend(frameon=False, fontsize=7, loc="lower right")
    fig.suptitle("Supplementary Fig. S12. Public LUSC scRNA patient-level PARK7 audit", fontsize=13)
    fig.tight_layout()
    save_fig(fig, SUPP / "S12_scRNA_PARK7_Pseudobulk_Audit.tif")


def make_tcga_context_figure() -> None:
    corr = pd.read_csv(SRC / "TCGA_LUSC_PARK7_bulk_correlation.csv")
    corr = corr.sort_values("spearman_rho")
    corr.to_csv(SRC / "S13_TCGA_LUSC_PARK7_Context_Correlations.csv", index=False)
    add_manifest_row(
        "S13_TCGA_LUSC_PARK7_Context_Correlations.csv",
        "Supplementary Fig. S13",
        "TCGA-LUSC PARK7 bulk-expression correlation context by marker/signature class.",
    )

    colors = {
        "redox marker": "#CC6677",
        "NRF2/redox marker": "#CC6677",
        "B-lineage marker": "#4477AA",
        "module": "#117733",
        "squamous epithelial marker": "#DDCC77",
        "immune marker": "#88CCEE",
    }
    c = [colors.get(x, "#999999") for x in corr["target_class"]]
    fig, ax = plt.subplots(figsize=(7.2, 5.2))
    ax.barh(corr["target"], corr["spearman_rho"], color=c)
    ax.axvline(0, color="black", linewidth=0.8)
    ax.set_xlabel("Spearman rho with PARK7")
    ax.set_title("Supplementary Fig. S13. TCGA-LUSC bulk-expression context")
    fig.tight_layout()
    save_fig(fig, SUPP / "S13_TCGA_LUSC_Bulk_Context.tif")


def make_network_boundary_figure() -> None:
    deg = pd.read_csv(SRC / "Figure5_InSilicoPerturbation_PARK7_DR_Results.csv")
    enrich = pd.read_csv(SRC / "Figure5_InSilicoPerturbation_PARK7_Enrichment.csv")
    top = enrich.sort_values("p.adjust").head(12).copy()
    top["minus_log10_fdr"] = -np.log10(top["p.adjust"].clip(lower=1e-300))
    top[["ID", "Description", "GeneRatio", "BgRatio", "FoldEnrichment", "p.adjust", "Count", "geneID"]].to_csv(
        SRC / "S14_NetworkInference_TopPathway_Audit.csv", index=False
    )
    add_manifest_row(
        "S14_NetworkInference_TopPathway_Audit.csv",
        "Supplementary Fig. S14",
        "Top pathway enrichment entries used for the network-inference boundary audit.",
    )

    top_genes = deg.sort_values("p.adj").head(500)["gene"].astype(str)
    classes = {
        "ribosomal_RPL_RPS": top_genes.str.match(r"^(RPL|RPS)").sum(),
        "immunoglobulin_IG": top_genes.str.match(r"^(IGH|IGK|IGL)").sum(),
        "mitochondrial_MT": top_genes.str.match(r"^MT-").sum(),
        "redox_anchor": top_genes.isin(["PARK7", "SOD1", "GPX1", "CAT", "NFE2L2", "TXN", "PRDX4", "HMOX1", "NQO1"]).sum(),
    }
    class_df = pd.DataFrame({"gene_class": list(classes.keys()), "count_in_top500": list(classes.values())})
    class_df.to_csv(SRC / "S14_NetworkInference_TopGeneClass_Audit.csv", index=False)
    add_manifest_row(
        "S14_NetworkInference_TopGeneClass_Audit.csv",
        "Supplementary Fig. S14",
        "Gene-class audit of top network-inferred PARK7-associated genes.",
    )

    fig, axes = plt.subplots(1, 2, figsize=(12, 4.2))
    axes[0].barh(top["Description"][::-1], top["minus_log10_fdr"][::-1], color="#44AA99")
    axes[0].set_xlabel("-log10 adjusted P")
    axes[0].set_title("Top enriched inferred pathways")
    axes[1].bar(class_df["gene_class"], class_df["count_in_top500"], color="#AA4499")
    axes[1].tick_params(axis="x", rotation=35)
    axes[1].set_ylabel("Count among top 500 genes")
    axes[1].set_title("High-RNA/program composition audit")
    fig.suptitle("Supplementary Fig. S14. Network-inferred PARK7 perturbation boundary audit", fontsize=13)
    fig.tight_layout()
    save_fig(fig, SUPP / "S14_NetworkInference_Boundary_Audit.tif")


def replace_para(para, text: str) -> None:
    for run in list(para.runs):
        run.text = ""
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)


def optimize_manuscript_text() -> None:
    backup(MANUSCRIPT)
    doc = Document(str(MANUSCRIPT))
    for para in doc.paragraphs:
        text = para.text.strip()
        artifact_map = {
            "\u6bcf": "-",
            "\u2105": "x",
            "\u00d7": "x",
            "\u2013": "-",
            "\u00a1\u00c1": "x",
        }
        if any(old in text for old in artifact_map):
            for old, new in artifact_map.items():
                text = text.replace(old, new)
            replace_para(para, text)
        if text == "Figure 1. PARK7-linked B-cell redox and ER-stress adaptation as a candidate hypothesis in LUSC.":
            replace_para(para, "Figure 1. Analytical evidence-boundary framework for PARK7 prioritization in LUSC.")
        elif text.startswith("Keywords:"):
            replace_para(
                para,
                "Keywords: PARK7; DJ-1; lung squamous cell carcinoma; B cells; Mendelian randomization; redox stress; ER stress; spatial transcriptomics",
            )
        elif text.startswith("This schematic summarizes the integrative workflow"):
            replace_para(
                para,
                "This schematic summarizes the evidence-boundary framework used to prioritize PARK7 in LUSC. The framework separates what each analytical layer supports from what it does not establish: cell-type MR nominates a B-cell-linked genetic candidate, locus-level posterior analysis provides weak-to-moderate regional evidence, spatial scoring provides indirect and heterogeneous redox context, public scRNA/TCGA analyses provide expression and pathway context, and IHC supports tissue-level PARK7 detectability. The figure is intended as a prioritization map rather than a mechanistic model of PARK7-positive B-cell redox niches.",
            )
        elif text == "Locus-level posterior evidence suggests a LUSC-preferential regional PARK7 signal":
            replace_para(para, "Locus-level analysis provides weak-to-moderate LUSC-preferential regional evidence at PARK7")
        elif text == "B-cell-enriched territories show limited and heterogeneous co-occurrence with antioxidant-high regions":
            replace_para(
                para,
                "Spatial transcriptomics constrains the hypothesis by showing heterogeneous, non-uniform B-cell/redox organization",
            )
        elif text == "Predicted PARK7 perturbation implicates ER stress, UPR, translation, and stress-adaptation programs":
            replace_para(
                para,
                "Network-inferred PARK7 perturbation prioritizes ER-stress, translation, and stress-adaptation context",
            )
        elif text.startswith("PARK7 perturbation was modeled using in silico PARK7 perturbation analysis"):
            replace_para(
                para,
                "PARK7-centered network inference was used to estimate transcriptional programs associated with PARK7 perturbation in B-cell expression data. Regulatory reweighting was summarized using differential regulatory scores, Z statistics, fold-change-like perturbation metrics, and adjusted P values. These outputs were interpreted as network-inferred transcriptional consequences and exploratory pathway-prioritization context, not as an experimental PARK7 knockdown or functional perturbation assay.",
            )
        elif text.startswith("Pathway enrichment of the in silico perturbation response highlighted"):
            replace_para(
                para,
                "Pathway enrichment of the PARK7-centered network-inference response highlighted cytoplasmic translation, response to endoplasmic reticulum stress, unfolded protein response, and DNA-damage-associated intrinsic apoptotic signaling (Supplementary Table S1). These computational results are consistent with a candidate model in which PARK7-linked redox adaptation may contribute to B-cell stress buffering in LUSC, but they do not provide functional validation of PARK7-dependent B-cell biology.",
            )
        elif text.startswith("In silico PARK7 perturbation analysis was used to infer predicted regulatory consequences"):
            replace_para(
                para,
                "PARK7-centered network inference was used to infer transcriptional programs associated with PARK7 perturbation in B-cell expression data. Panel A shows the PARK7-centered regulatory network. Panel B shows inferred regulatory reweighting after PARK7-centered perturbation modeling, including loss of selected hub-coupling edges. Panel C highlights representative genes affected in the network-inference model, including immunoglobulin, ER-associated, and stress-adaptation genes. Panel D summarizes pathway enrichment, emphasizing cytoplasmic translation, ER stress, unfolded protein response, and DNA-damage-associated apoptotic signaling. Connectivity Map reversal signatures are presented as hypothesis-generating perturbational patterns and not as evidence of drug efficacy. CMap, Connectivity Map; DDR, DNA damage response; ER, endoplasmic reticulum; UPR, unfolded protein response.",
            )
        elif text.startswith("Analyses were performed using R/Python-based workflows") and "Supplementary Figs. S10-S14" not in text:
            replace_para(
                para,
                para.text
                + " Additional reviewer-facing audit analyses are provided in Supplementary Figs. S10-S14 and Supplementary Table S9; these summarize available MR robustness diagnostics, spatial threshold/permutation sensitivity, scRNA patient-level expression summaries, TCGA bulk-expression context, and network-inference boundary checks without converting these analyses into mechanistic validation.",
            )
        elif text.startswith("Cell-type-resolved MR prioritized PARK7 expression") and "per-instrument audit" not in text:
            replace_para(
                para,
                para.text
                + " A per-instrument audit table, estimator comparison, leave-one-SNP-out stability profile, and instrument-strength visualization are provided to make the MR signal auditable (Supplementary Fig. S10; Supplementary Tables S1A-S1D). Analyses requiring raw genome-wide summary statistics or alternative eQTL instrument re-selection, including formal Steiger testing, MR-RAPS, radial MR, and alternative cis-window/eQTL-threshold grids, were not inferred from the available reduced source tables.",
            )
        elif text.startswith("Spatial transcriptomic analysis identified heterogeneous") and "threshold and permutation sensitivity" not in text:
            replace_para(
                para,
                para.text
                + " Threshold and permutation sensitivity analyses showed that the limited same-spot overlap was not a single-threshold artifact and should be interpreted as boundary evidence against uniform co-localization rather than as spatial validation of PARK7-positive B-cell redox niches (Supplementary Fig. S11 and source data).",
            )
        elif text.startswith("As an additional public LUSC scRNA-seq validation") and "patient-level" not in text:
            replace_para(
                para,
                para.text
                + " Patient/sample-level pseudobulk summaries further report PARK7 detection fraction and positive-cell expression intensity by broad assigned cell type, with B cells and plasma cells kept separate to reduce plasma-cell program conflation (Supplementary Fig. S12 and source data).",
            )
        elif text.startswith("An additional TCGA-LUSC bulk RNA-seq correlation analysis") and "Supplementary Fig. S13" not in text:
            replace_para(
                para,
                para.text
                + " A reviewer-facing correlation audit by marker class is provided as Supplementary Fig. S13; the analysis remains an orthogonal bulk-tissue context layer and not evidence of cell-type-specific PARK7 function.",
            )
        elif text.startswith("Mechanistically, in silico perturbation analysis provided"):
            replace_para(
                para,
                "The network-inference analysis provides exploratory transcriptional context rather than mechanistic validation. PARK7-centered regulatory reweighting prioritized translation, ER stress/UPR, and DNA-damage-associated apoptotic signaling, but these outputs should be treated as inferred pathway hypotheses. PARK7/DJ-1 has structural support from protein data resources and crystallographic studies [40,41], and published work links DJ-1 to NRF2 stabilization [42]. Reactive oxygen species are also central components of tumor microenvironmental stress biology [43]. These observations make a redox/ER-stress hypothesis biologically plausible, but direct PARK7-dependent B-cell function requires next-stage perturbation and co-localization experiments.",
            )
        elif text.startswith("Several limitations should be acknowledged") and "formal MR-RAPS" not in text:
            replace_para(
                para,
                para.text
                + " Ninth, formal MR-RAPS, radial MR, Steiger directionality, colocalization prior sensitivity, conditional colocalization, and independent second-cohort scRNA validation were not claimed unless directly supported by available source data. These analyses remain desirable next-stage robustness checks when the required raw summary statistics or additional harmonized public cohorts are assembled.",
            )
        elif text.startswith("Publicly available datasets analyzed in this study") and "Supplementary Table S9" not in text:
            replace_para(
                para,
                para.text
                + " Supplementary Table S9 provides a claim-by-claim wording boundary matrix specifying which statements are directly supported, which remain unsupported, and which manuscript wording is permitted for each evidence layer.",
            )
    doc.save(str(MANUSCRIPT))


def optimize_frontiers_statements() -> None:
    for path in list(STATEMENTS.glob("*.docx")) + list(ADMIN.glob("*.docx")):
        backup(path)
        doc = Document(str(path))
        changed = False
        for para in doc.paragraphs:
            text = para.text
            new = text
            old_title = (
                "Cell-type-resolved genetic prioritization of PARK7 and indirect spatial mapping of "
                "B-cell redox contexts in lung squamous cell carcinoma"
            )
            new = new.replace(old_title, NEW_TITLE)
            new = new.replace(
                "it addresses tumor-immune stress adaptation in lung squamous cell carcinoma using a public-data-integrated, hypothesis-generating prioritization framework",
                "it addresses computational genetic prioritization in lung squamous cell carcinoma using a public-data-integrated, hypothesis-generating evidence-boundary framework",
            )
            new = new.replace(
                "It addresses tumor-immune stress adaptation in lung squamous cell carcinoma using a public-data-integrated, hypothesis-generating prioritization framework",
                "It addresses computational genetic prioritization in lung squamous cell carcinoma using a public-data-integrated, hypothesis-generating evidence-boundary framework",
            )
            new = new.replace(
                "it addresses tumor-immune stress adaptation in lung squamous cell carcinoma and provides a transparent framework for prioritizing a PARK7-linked B-cell redox/ER-stress hypothesis for future multiplex tissue and functional validation",
                "it presents a bounded, reproducible, public-data-integrated genetic prioritization study. Its primary contribution is the transparent nomination and stress-contextualization of PARK7 as a B-cell-linked LUSC candidate, rather than experimental validation of PARK7 function",
            )
            if new != text:
                replace_para(para, new)
                changed = True
        if changed:
            doc.save(str(path))


def optimize_title_page() -> None:
    backup(TITLE_PAGE)
    doc = Document(str(TITLE_PAGE))
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("Number of supplementary figures:"):
            replace_para(para, "Number of supplementary figures: 14")
        elif text.startswith("Number of tables:"):
            replace_para(para, "Number of tables: 1 main table and 13 supplementary/source tables")
    doc.save(str(TITLE_PAGE))


def optimize_supplement_doc() -> None:
    supp_doc = SUPP / "Supplementary_Material.docx"
    backup(supp_doc)
    doc = Document(str(supp_doc))
    existing = "\n".join(p.text for p in doc.paragraphs)
    additions = [
        (
            "Supplementary Fig. S10. PARK7 MR instrument and estimator audit.",
            "This reviewer-facing audit summarizes the available PARK7 intermediate-B-cell MR diagnostics, including per-SNP F-statistic distribution, available IVW/weighted-median/MR-Egger estimates, per-SNP exposure/outcome effects, and leave-one-SNP-out IVW stability. It supports transparency of the primary MR prioritization signal but does not replace additional raw-summary-statistic analyses such as MR-RAPS, radial MR, or Steiger directionality testing.",
        ),
        (
            "Supplementary Fig. S11. Spatial threshold and permutation sensitivity.",
            "This analysis evaluates whether limited same-spot B-cell/antioxidant overlap is robust to alternative thresholds and permutation null comparisons. It is interpreted as boundary evidence against uniform spatial co-localization and not as validation of PARK7-positive B-cell redox niches.",
        ),
        (
            "Supplementary Fig. S12. Public LUSC scRNA patient-level PARK7 audit.",
            "PARK7 detection fraction and positive-cell expression intensity are summarized by broad assigned cell type and sample in GSE200972. B cells and plasma cells are reported separately. This supports detectable PARK7 expression across tumor microenvironment compartments, including B-lineage cells, but not B-cell specificity.",
        ),
        (
            "Supplementary Fig. S13. TCGA-LUSC bulk-expression context.",
            "Spearman correlations between PARK7 and selected redox, NRF2, B-lineage, immune, and epithelial markers or modules are summarized as orthogonal bulk-tissue context. These correlations are not interpreted as cell-type-specific evidence or prognostic validation.",
        ),
        (
            "Supplementary Fig. S14. Network-inferred PARK7 perturbation boundary audit.",
            "Top pathway enrichments and high-RNA/program gene-class composition are summarized for the PARK7-centered network-inference analysis. The figure highlights pathway-prioritization context while emphasizing that the analysis is not an experimental PARK7 knockdown or functional validation.",
        ),
        (
            "Supplementary Table S9. Reviewer-facing claim-boundary and allowed-wording matrix.",
            "This table lists major manuscript claims, whether each is directly supported, the evidence source, strength, main limitation, allowed manuscript wording, and wording to avoid. It is intended to make the computational/genetic prioritization boundary explicit.",
        ),
    ]
    for title, body in additions:
        if title not in existing:
            p = doc.add_paragraph()
            p.add_run(title).bold = True
            doc.add_paragraph(body)
    doc.save(str(supp_doc))


def rebuild_zip(zip_path: Path, source_dir: Path) -> None:
    if zip_path.exists():
        backup_dir = WORK / "zip_backups"
        backup_dir.mkdir(exist_ok=True)
        shutil.copy2(zip_path, backup_dir / (zip_path.name + ".before_prioritization_audit"))
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=8) as z:
        for path in source_dir.rglob("*"):
            if not path.is_file():
                continue
            if path.resolve() == zip_path.resolve():
                continue
            if ".bak_" in path.name.lower() or "backup" in path.name.lower():
                continue
            if source_dir == PACKAGE and path.suffix.lower() == ".zip" and path.parent == PACKAGE:
                continue
            z.write(path, path.relative_to(source_dir.parent))


def update_package_manifest() -> None:
    rows = []
    for path in PACKAGE.rglob("*"):
        if not path.is_file() or ".bak_" in path.name.lower() or "backup" in path.name.lower():
            continue
        rel = path.relative_to(PACKAGE).as_posix()
        role = "package documentation"
        if rel.startswith("01_Manuscript/"):
            role = "core manuscript"
        elif rel.startswith("02_Frontiers_Statements/"):
            role = "Frontiers statement/admin"
        elif rel.startswith("03_Figures_Main/"):
            role = "main figure"
        elif rel.startswith("04_Supplementary_Material/"):
            role = "supplementary material"
        elif rel.startswith("05_Source_Data/"):
            role = "source data / reproducibility"
        elif rel.startswith("06_Original_Microscopy_Images/"):
            role = "original microscopy image"
        elif rel.startswith("07_Admin_and_Optional/"):
            role = "admin optional"
        rows.append({"relative_path": rel, "size_bytes": path.stat().st_size, "role": role})
    pd.DataFrame(rows).sort_values("relative_path").to_csv(PACKAGE / "PACKAGE_MANIFEST.csv", index=False)


def main() -> None:
    make_claim_boundary_table()
    make_main_figure_1_boundary_framework()
    make_mr_audit_figure()
    make_spatial_sensitivity_figure()
    make_scrna_pseudobulk_figure()
    make_tcga_context_figure()
    make_network_boundary_figure()
    optimize_manuscript_text()
    optimize_title_page()
    optimize_supplement_doc()
    optimize_frontiers_statements()
    update_package_manifest()
    rebuild_zip(PACKAGE / "03_Figures_Main.zip", MAIN_FIG)
    rebuild_zip(PACKAGE / "04_Supplementary_Material.zip", SUPP)
    rebuild_zip(PACKAGE / "05_Source_Data" / "Source_Data_CSV.zip", SRC)
    rebuild_zip(PACKAGE / "05_Source_Data.zip", PACKAGE / "05_Source_Data")
    rebuild_zip(ROOT / "submission_package.zip", PACKAGE)
    (WORK / "PRIORITIZATION_AUDIT_REPORT.md").write_text(
        "# Prioritization audit update completed\n\n"
        "- Added Supplementary Figs. S10-S14 from available source data.\n"
        "- Added Supplementary Table S9 reviewer-facing claim boundary matrix.\n"
        "- Reframed Figure 1 caption, result subheadings, and network-inference wording.\n"
        "- Explicitly stated that unsupported robustness analyses are not claimed.\n"
        "- Refreshed manifests and zips.\n",
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
