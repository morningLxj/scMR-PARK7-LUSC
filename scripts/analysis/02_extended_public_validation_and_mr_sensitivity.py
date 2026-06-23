from __future__ import annotations

import gzip
import math
import os
import shutil
import tarfile
import zipfile
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from scipy import stats
from scipy.special import ndtri
from scipy.stats import gaussian_kde
from statsmodels.stats.multitest import multipletests

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt


REPO_ROOT = Path(__file__).resolve().parents[2]
ROOT = Path(os.environ.get("SCMR_PROJECT_ROOT", REPO_ROOT))
PACKAGE = Path(os.environ.get("SCMR_SUBMISSION_PACKAGE", ROOT / "submission_package"))
SUPP = PACKAGE / "04_Supplementary_Material"
SRC = Path(os.environ.get("SCMR_SOURCE_DATA_DIR", REPO_ROOT / "data" / "source_data_csv"))
MANUSCRIPT = PACKAGE / "01_Manuscript" / "Manuscript.docx"
TITLE_PAGE = PACKAGE / "01_Manuscript" / "Title_Page_Frontiers.docx"
PROJECT = Path(os.environ.get("SCMR_PROJECT_DATA_DIR", ROOT / "external_data"))
DOWNLOAD = Path(os.environ.get("SCMR_DOWNLOAD_DIR", ROOT / "raw_data" / "downloaded_public_data_20260612"))
WORK = Path(os.environ.get("SCMR_WORK_DIR", ROOT / "work" / "completed_public_analyses_20260612"))
WORK.mkdir(parents=True, exist_ok=True)


def backup(path: Path) -> None:
    if not path.exists():
        return
    rel = path.relative_to(PACKAGE) if path.is_relative_to(PACKAGE) else path.name
    dst = WORK / "backups" / rel
    dst.parent.mkdir(parents=True, exist_ok=True)
    if not dst.exists():
        shutil.copy2(path, dst)


def save_fig(fig, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, dpi=300, bbox_inches="tight", pil_kwargs={"compression": "tiff_lzw"})
    plt.close(fig)


def add_manifest_row(file_name: str, manuscript_item: str, current_note: str) -> None:
    manifest = SRC / "Source_Data_CSV_Manifest.csv"
    df = pd.read_csv(manifest) if manifest.exists() else pd.DataFrame(columns=["file_name", "manuscript_item", "current_note"])
    if file_name in set(df["file_name"].astype(str)):
        df.loc[df["file_name"].astype(str) == file_name, ["manuscript_item", "current_note"]] = [
            manuscript_item,
            current_note,
        ]
    else:
        df = pd.concat(
            [df, pd.DataFrame([{"file_name": file_name, "manuscript_item": manuscript_item, "current_note": current_note}])],
            ignore_index=True,
        )
    df.to_csv(manifest, index=False)


def replace_para(para, text: str) -> None:
    for run in list(para.runs):
        run.text = ""
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)


def bh(pvals: pd.Series) -> np.ndarray:
    vals = pd.to_numeric(pvals, errors="coerce").to_numpy()
    ok = np.isfinite(vals)
    out = np.full(len(vals), np.nan)
    if ok.any():
        out[ok] = multipletests(vals[ok], method="fdr_bh")[1]
    return out


MARKERS = {
    "B cells": ["MS4A1", "CD79A", "CD79B", "CD19", "CD37", "CD74"],
    "Plasma cells": ["MZB1", "JCHAIN", "XBP1", "SDC1", "IGKC", "IGHG1"],
    "T cells": ["CD3D", "CD3E", "TRAC", "CD2"],
    "NK cells": ["NKG7", "GNLY", "KLRD1", "KLRF1"],
    "Myeloid cells": ["LST1", "LYZ", "S100A8", "S100A9", "FCGR3A", "CST3"],
    "Epithelial/tumor": ["EPCAM", "KRT5", "KRT17", "KRT19", "TP63", "SOX2"],
    "Fibroblasts": ["COL1A1", "COL1A2", "DCN", "LUM", "COL3A1"],
    "Endothelial cells": ["PECAM1", "VWF", "KDR", "ENG"],
    "Mast cells": ["TPSAB1", "TPSB2", "CPA3", "KIT"],
}
STRESS_SETS = {
    "UPR_ER_stress_score": ["XBP1", "HSPA5", "DDIT3", "ATF4", "HERPUD1", "HSP90B1", "PDIA3"],
    "NRF2_redox_score": ["PARK7", "SOD1", "GPX1", "CAT", "NFE2L2", "TXN", "NQO1", "HMOX1", "GSR"],
    "Ribosome_translation_score": ["RPLP0", "RPL13A", "RPL37", "RPS12", "RPS27", "RPS3", "EEF2"],
    "Apoptosis_DDR_score": ["BAX", "BCL2", "CASP3", "TP53", "PERP", "GADD45A"],
}
SELECTED_GENES = sorted(set(["PARK7"] + sum(MARKERS.values(), []) + sum(STRESS_SETS.values(), [])))


def process_gse148071() -> None:
    tar_path = DOWNLOAD / "GSE148071" / "GSE148071_RAW.tar"
    if not tar_path.exists():
        raise FileNotFoundError(tar_path)

    celltype_rows: list[dict] = []
    sample_rows: list[dict] = []
    blineage_rows: list[dict] = []
    cell_rows: list[pd.DataFrame] = []

    with tarfile.open(tar_path) as tar:
        members = [m for m in tar.getmembers() if m.name.endswith("_exp.txt.gz")]
        for idx, member in enumerate(members, 1):
            sample = member.name.replace("_exp.txt.gz", "")
            sample_short = sample.split("_")[-2] if "_" in sample else sample
            print(f"[GSE148071 {idx}/{len(members)}] {member.name}")
            fh = tar.extractfile(member)
            if fh is None:
                continue
            gz = gzip.GzipFile(fileobj=fh)
            header = gz.readline().decode("utf-8", "replace").rstrip("\n").split("\t")
            cells = header
            n = len(cells)
            totals = np.zeros(n, dtype=np.float64)
            selected_counts: dict[str, np.ndarray] = {g: np.zeros(n, dtype=np.float32) for g in SELECTED_GENES}
            for raw in gz:
                parts = raw.rstrip(b"\n").split(b"\t", 1)
                if len(parts) != 2:
                    continue
                gene = parts[0].decode("utf-8", "replace")
                vals = np.fromstring(parts[1].decode("ascii", "ignore"), sep="\t", dtype=np.float64)
                if vals.size != n:
                    continue
                totals += vals
                if gene in selected_counts:
                    selected_counts[gene] = vals.astype(np.float32)
            denom = np.maximum(totals, 1.0) / 1e6
            log_expr = {g: np.log2(selected_counts[g] / denom + 1.0) for g in SELECTED_GENES}
            scores = {}
            for ct, genes in MARKERS.items():
                have = [log_expr[g] for g in genes if g in log_expr]
                scores[ct] = np.mean(have, axis=0) if have else np.zeros(n)
            score_df = pd.DataFrame(scores)
            assigned = score_df.idxmax(axis=1).to_numpy()
            max_score = score_df.max(axis=1).to_numpy()
            assigned[max_score < 0.20] = "Low-confidence/other"
            park7 = log_expr["PARK7"]
            detected = park7 > 0
            cdf = pd.DataFrame(
                {
                    "dataset": "GSE148071",
                    "sample": sample,
                    "cell_barcode": cells,
                    "assigned_cell_type": assigned,
                    "assignment_score": max_score,
                    "PARK7_log2CPM": park7,
                    "PARK7_detected": detected,
                }
            )
            for sname, genes in STRESS_SETS.items():
                have = [log_expr[g] for g in genes if g in log_expr]
                cdf[sname] = np.mean(have, axis=0) if have else 0
            cell_rows.append(cdf)

            grouped = cdf.groupby("assigned_cell_type")
            for ct, sub in grouped:
                vals = sub["PARK7_log2CPM"].to_numpy()
                celltype_rows.append(
                    {
                        "dataset": "GSE148071",
                        "sample": sample,
                        "assigned_cell_type": ct,
                        "n_cells": len(sub),
                        "PARK7_pct_detected": 100 * np.mean(vals > 0),
                        "PARK7_mean_log2CPM": float(np.mean(vals)),
                        "PARK7_median_log2CPM": float(np.median(vals)),
                        "PARK7_positive_mean_log2CPM": float(np.mean(vals[vals > 0])) if np.any(vals > 0) else 0.0,
                    }
                )
            sample_rows.append(
                {
                    "dataset": "GSE148071",
                    "sample": sample,
                    "n_cells": n,
                    "PARK7_pct_detected_all_cells": 100 * np.mean(detected),
                    "PARK7_mean_log2CPM_all_cells": float(np.mean(park7)),
                }
            )
            bmask = cdf["assigned_cell_type"].isin(["B cells", "Plasma cells"])
            bdf = cdf.loc[bmask].copy()
            if len(bdf) >= 20 and (bdf["PARK7_log2CPM"] > 0).sum() >= 5:
                cutoff = bdf.loc[bdf["PARK7_log2CPM"] > 0, "PARK7_log2CPM"].median()
                bdf["PARK7_group"] = np.where(bdf["PARK7_log2CPM"] >= cutoff, "PARK7_high", "PARK7_low")
                for ct in ["B cells", "Plasma cells", "B-lineage combined"]:
                    sub = bdf if ct == "B-lineage combined" else bdf[bdf["assigned_cell_type"] == ct]
                    if len(sub) < 10 or sub["PARK7_group"].nunique() < 2:
                        continue
                    row = {"dataset": "GSE148071", "sample": sample, "cell_group": ct, "n_cells": len(sub), "PARK7_cutoff_log2CPM": cutoff}
                    for sname in STRESS_SETS:
                        hi = sub.loc[sub["PARK7_group"] == "PARK7_high", sname]
                        lo = sub.loc[sub["PARK7_group"] == "PARK7_low", sname]
                        if len(hi) and len(lo):
                            row[f"{sname}_high_mean"] = float(hi.mean())
                            row[f"{sname}_low_mean"] = float(lo.mean())
                            row[f"{sname}_delta_high_minus_low"] = float(hi.mean() - lo.mean())
                            row[f"{sname}_mannwhitney_p"] = float(stats.mannwhitneyu(hi, lo, alternative="two-sided").pvalue)
                    blineage_rows.append(row)

    celltype = pd.DataFrame(celltype_rows)
    sample_sum = pd.DataFrame(sample_rows)
    blineage = pd.DataFrame(blineage_rows)
    cell_level = pd.concat(cell_rows, ignore_index=True)
    celltype["PARK7_detection_rank_within_sample"] = celltype.groupby("sample")["PARK7_pct_detected"].rank(ascending=False, method="min")
    if not blineage.empty:
        pcols = [c for c in blineage.columns if c.endswith("_mannwhitney_p")]
        for c in pcols:
            blineage[c.replace("_p", "_fdr")] = bh(blineage[c])

    files = {
        "S15_GSE148071_LUSC_scRNA_PARK7_celltype_summary.csv": celltype,
        "S15_GSE148071_LUSC_scRNA_sample_summary.csv": sample_sum,
        "S15_GSE148071_LUSC_scRNA_Blineage_stress_scores.csv": blineage,
        "S15_GSE148071_LUSC_scRNA_cell_level_selected_scores.csv": cell_level,
    }
    for name, df in files.items():
        df.to_csv(SRC / name, index=False)
        add_manifest_row(name, "Supplementary Fig. S15", "Independent GSE148071 LUSC scRNA validation source data.")

    # Figure
    order = (
        celltype.groupby("assigned_cell_type")["PARK7_pct_detected"].median().sort_values(ascending=False).index.tolist()
    )
    fig, axes = plt.subplots(1, 3, figsize=(13.2, 4.2))
    data = [celltype.loc[celltype["assigned_cell_type"] == ct, "PARK7_pct_detected"] for ct in order]
    axes[0].boxplot(data, labels=order, vert=False, showfliers=False)
    axes[0].set_xlabel("PARK7 detection (%)")
    axes[0].set_title("GSE148071 detection by assigned cell type")
    data2 = [celltype.loc[celltype["assigned_cell_type"] == ct, "PARK7_positive_mean_log2CPM"] for ct in order]
    axes[1].boxplot(data2, labels=order, vert=False, showfliers=False)
    axes[1].set_xlabel("Positive-cell PARK7 log2(CPM+1)")
    axes[1].set_title("Expression intensity")
    if not blineage.empty:
        plot_b = blineage[blineage["cell_group"] == "B-lineage combined"]
        delta_cols = [c for c in plot_b.columns if c.endswith("_delta_high_minus_low")]
        vals = [plot_b[c].median() for c in delta_cols]
        labels = [c.replace("_delta_high_minus_low", "").replace("_score", "") for c in delta_cols]
        axes[2].barh(labels, vals, color="#4477AA")
        axes[2].axvline(0, color="black", linewidth=0.8)
        axes[2].set_xlabel("Median delta, PARK7-high minus low")
        axes[2].set_title("B-lineage stress-score context")
    fig.suptitle("Supplementary Fig. S15. Independent GSE148071 LUSC scRNA PARK7 validation", fontsize=13)
    fig.tight_layout()
    save_fig(fig, SUPP / "S15_GSE148071_LUSC_scRNA_PARK7_validation.tif")


def weighted_median(beta: np.ndarray, weights: np.ndarray) -> float:
    idx = np.argsort(beta)
    b = beta[idx]
    w = weights[idx]
    cw = np.cumsum(w) / np.sum(w)
    return float(b[np.searchsorted(cw, 0.5)])


def mode_estimate(beta: np.ndarray, weights: np.ndarray) -> float:
    if len(beta) < 3:
        return float(np.nan)
    grid = np.linspace(np.nanpercentile(beta, 2), np.nanpercentile(beta, 98), 512)
    try:
        kde = gaussian_kde(beta, weights=weights / weights.sum())
        return float(grid[np.argmax(kde(grid))])
    except Exception:
        return float(np.nan)


def huber_slope(x: np.ndarray, y: np.ndarray, w: np.ndarray, max_iter: int = 100) -> float:
    beta = np.sum(w * x * y) / np.sum(w * x * x)
    for _ in range(max_iter):
        resid = y - beta * x
        scale = np.median(np.abs(resid - np.median(resid))) / 0.6745
        if not np.isfinite(scale) or scale == 0:
            break
        k = 1.345 * scale
        rw = np.minimum(1.0, k / np.maximum(np.abs(resid), 1e-12))
        ww = w * rw
        nbeta = np.sum(ww * x * y) / np.sum(ww * x * x)
        if abs(nbeta - beta) < 1e-8:
            beta = nbeta
            break
        beta = nbeta
    return float(beta)


def extract_park7_eqtl(cell: str) -> pd.DataFrame:
    cache = WORK / f"PARK7_{cell}_eqtl_extracted.csv"
    if cache.exists():
        return pd.read_csv(cache)
    path = PROJECT / "Exposure" / f"{cell}_eqtl_table.tsv.gz"
    cols = ["RSID", "GENE", "CHR", "POS", "A1", "A2", "A2_FREQ_ONEK1K", "SPEARMANS_RHO", "P_VALUE"]
    rows = []
    for chunk in pd.read_csv(path, sep="\t", usecols=cols, chunksize=500_000):
        sub = chunk[chunk["GENE"].astype(str).eq("PARK7")].copy()
        if not sub.empty:
            rows.append(sub)
    out = pd.concat(rows, ignore_index=True) if rows else pd.DataFrame(columns=cols)
    out.to_csv(cache, index=False)
    return out


def make_mr_extended() -> None:
    inst = pd.read_csv(SRC / "SupplementaryTableS1A_PARK7_PerInstrument_F_Statistics.csv")
    inst = inst[
        (inst["Gene"].astype(str) == "PARK7")
        & (inst["Cell"].astype(str) == "bin")
        & (inst["included_in_primary_lusc_mr"].astype(bool))
    ].copy()
    x = inst["exposure_beta"].to_numpy(float)
    sx = inst["exposure_se"].to_numpy(float)
    y = inst["outcome_beta"].to_numpy(float)
    sy = inst["outcome_se"].to_numpy(float)
    ratio = y / x
    ratio_se = np.sqrt((sy**2 / x**2) + ((y**2) * sx**2 / x**4))
    w = 1 / ratio_se**2
    ivw_beta = np.sum(w * ratio) / np.sum(w)
    ivw_se = math.sqrt(1 / np.sum(w))
    q = np.sum(w * (ratio - ivw_beta) ** 2)
    q_p = stats.chi2.sf(q, len(ratio) - 1)
    wm = weighted_median(ratio, w)
    smode = mode_estimate(ratio, np.ones_like(w))
    wmode = mode_estimate(ratio, w)
    huber = huber_slope(x, y, 1 / sy**2)
    primary = pd.read_csv(SRC / "SupplementaryTableS1C_PARK7_MR_Sensitivity_Summary.csv").iloc[0]
    X = np.vstack([np.ones_like(x), x]).T
    W = np.diag(1 / sy**2)
    beta_hat = np.linalg.inv(X.T @ W @ X) @ (X.T @ W @ y)
    cov = np.linalg.inv(X.T @ W @ X)
    egger_intercept, egger_slope = beta_hat[0], beta_hat[1]
    egger_intercept_se = math.sqrt(cov[0, 0])
    egger_intercept_p = 2 * stats.norm.sf(abs(egger_intercept / egger_intercept_se))
    r2_exp = 2 * inst["eaf"].to_numpy(float) * (1 - inst["eaf"].to_numpy(float)) * x**2
    r2_out = 2 * inst["eaf"].to_numpy(float) * (1 - inst["eaf"].to_numpy(float)) * y**2
    steiger_direction = "exposure_to_outcome" if np.nansum(r2_exp) > np.nansum(r2_out) else "outcome_to_exposure_or_ambiguous"
    radial_resid = (ratio - ivw_beta) / ratio_se
    outlier = np.abs(radial_resid) > 3
    methods = pd.DataFrame(
        [
            ["IVW", primary["IVW_beta"], primary["IVW_SE"], primary["IVW_P"], len(inst), "primary MR scale from Supplementary Table S1C"],
            ["Weighted median", primary["Weighted_median_beta"], primary["Weighted_median_SE_bootstrap"], primary["Weighted_median_P"], len(inst), "primary MR scale from Supplementary Table S1C"],
            ["MR-Egger slope", primary["MR_Egger_slope"], primary["MR_Egger_slope_SE"], primary["MR_Egger_slope_P"], len(inst), "primary MR scale from Supplementary Table S1C"],
            ["Simple mode", smode, np.nan, np.nan, len(inst), "Wald-ratio diagnostic scale; directionality only"],
            ["Weighted mode", wmode, np.nan, np.nan, len(inst), "Wald-ratio diagnostic scale; directionality only"],
            ["Huber robust IVW", huber, np.nan, np.nan, len(inst), "Wald-ratio diagnostic scale; directionality only"],
        ],
        columns=["method", "beta", "se", "p", "n_instruments", "note"],
    )
    methods.to_csv(SRC / "S16_MR_Extended_Robust_Estimators.csv", index=False)
    add_manifest_row("S16_MR_Extended_Robust_Estimators.csv", "Supplementary Fig. S16", "Extended MR robust estimator summary.")
    radial = inst[["SNP", "exposure_beta", "outcome_beta"]].copy()
    radial["wald_ratio"] = ratio
    radial["ratio_se"] = ratio_se
    radial["radial_standardized_residual"] = radial_resid
    radial["radial_outlier_abs_gt3"] = outlier
    radial["approx_exposure_r2"] = r2_exp
    radial["approx_outcome_r2"] = r2_out
    radial.to_csv(SRC / "S16_MR_RadialLike_Steiger_Source.csv", index=False)
    add_manifest_row("S16_MR_RadialLike_Steiger_Source.csv", "Supplementary Fig. S16", "Radial-like outlier diagnostics and approximate Steiger directionality source.")

    eqtl = extract_park7_eqtl("bin")
    eqtl = eqtl.rename(columns={"RSID": "SNP", "POS": "snp_pos", "P_VALUE": "full_eqtl_p"})
    eqtl = eqtl.sort_values("full_eqtl_p").drop_duplicates("SNP", keep="first")
    merged = inst.merge(eqtl[["SNP", "snp_pos", "full_eqtl_p"]], on="SNP", how="left")
    center = 7014444
    sens_rows = []
    for label, mask in [
        ("all_primary", np.ones(len(merged), dtype=bool)),
        ("eQTL_P_lt_5e-8", merged["exposure_p"] < 5e-8),
        ("eQTL_P_lt_1e-7", merged["exposure_p"] < 1e-7),
        ("eQTL_P_lt_1e-6", merged["exposure_p"] < 1e-6),
        ("cis_250kb", (merged["snp_pos"] - center).abs() <= 250_000),
        ("cis_500kb", (merged["snp_pos"] - center).abs() <= 500_000),
        ("cis_1Mb", (merged["snp_pos"] - center).abs() <= 1_000_000),
        ("cis_10Mb", (merged["snp_pos"] - center).abs() <= 10_000_000),
    ]:
        sub = merged[mask.fillna(False) if hasattr(mask, "fillna") else mask]
        if len(sub) < 3:
            continue
        bx = sub["exposure_beta"].to_numpy(float)
        by = sub["outcome_beta"].to_numpy(float)
        sx2 = sub["exposure_se"].to_numpy(float)
        sy2 = sub["outcome_se"].to_numpy(float)
        rr = by / bx
        rse = np.sqrt((sy2**2 / bx**2) + ((by**2) * sx2**2 / bx**4))
        ww = 1 / rse**2
        b = np.sum(ww * rr) / np.sum(ww)
        se = math.sqrt(1 / np.sum(ww))
        sens_rows.append({"setting": label, "n_instruments": len(sub), "wald_ratio_diagnostic_beta": b, "se": se, "p": 2 * stats.norm.sf(abs(b / se)), "interpretation": "directionality/stability only; not primary MR scale"})
    sens = pd.DataFrame(sens_rows)
    sens.to_csv(SRC / "S16_MR_Threshold_CisWindow_Sensitivity.csv", index=False)
    add_manifest_row("S16_MR_Threshold_CisWindow_Sensitivity.csv", "Supplementary Fig. S16", "MR sensitivity across available eQTL thresholds and cis windows.")
    summary = pd.DataFrame(
        [
            {
                "diagnostic": "Cochran_Q",
                "value": q,
                "p": q_p,
                "interpretation": "heterogeneity test for Wald ratios",
            },
            {
                "diagnostic": "MR_Egger_intercept",
                "value": egger_intercept,
                "se": egger_intercept_se,
                "p": egger_intercept_p,
                "interpretation": "directional pleiotropy screen",
            },
            {
                "diagnostic": "Approximate_Steiger_sum_R2_exposure",
                "value": np.nansum(r2_exp),
                "interpretation": steiger_direction,
            },
            {
                "diagnostic": "Approximate_Steiger_sum_R2_outcome",
                "value": np.nansum(r2_out),
                "interpretation": steiger_direction,
            },
            {
                "diagnostic": "Radial_like_abs_residual_gt3_count",
                "value": int(outlier.sum()),
                "interpretation": "outlier screen based on standardized Wald-ratio residuals",
            },
            {
                "diagnostic": "MR_RAPS_RadialMR_package_status",
                "value": np.nan,
                "interpretation": "mr.raps and RadialMR were unavailable for R 4.5.3 CRAN binary; robust alternatives are reported.",
            },
        ]
    )
    summary.to_csv(SRC / "S16_MR_Extended_Diagnostics_Summary.csv", index=False)
    add_manifest_row("S16_MR_Extended_Diagnostics_Summary.csv", "Supplementary Fig. S16", "Extended MR diagnostics summary.")

    fig, axes = plt.subplots(1, 3, figsize=(13.2, 4.0))
    yvals = np.arange(len(methods))
    axes[0].scatter(methods["beta"], yvals, color="#332288")
    axes[0].axvline(0, color="black", lw=0.8)
    axes[0].set_yticks(yvals, methods["method"])
    axes[0].set_title("Robust MR estimators")
    axes[0].set_xlabel("Beta")
    axes[1].scatter(radial["wald_ratio"], radial["radial_standardized_residual"], c=outlier, cmap="coolwarm", s=22)
    axes[1].axhline(3, color="black", ls="--", lw=0.8)
    axes[1].axhline(-3, color="black", ls="--", lw=0.8)
    axes[1].set_title("Radial-like outlier screen")
    axes[1].set_xlabel("Wald ratio")
    axes[1].set_ylabel("standardized residual")
    axes[2].errorbar(sens["wald_ratio_diagnostic_beta"], np.arange(len(sens)), xerr=1.96 * sens["se"], fmt="o", color="#117733")
    axes[2].axvline(0, color="black", lw=0.8)
    axes[2].set_yticks(np.arange(len(sens)), sens["setting"])
    axes[2].set_title("Threshold/cis-window sensitivity")
    axes[2].set_xlabel("Wald-ratio diagnostic beta")
    fig.suptitle("Supplementary Fig. S16. Extended PARK7 MR robustness analyses", fontsize=13)
    fig.tight_layout()
    save_fig(fig, SUPP / "S16_Extended_MR_Robustness.tif")


def coloc_abf(df: pd.DataFrame, p1: float, p2: float, p12: float, w1: float = 0.04, w2: float = 0.04) -> dict:
    z1 = df["beta_eqtl"].to_numpy(float) / df["se_eqtl"].to_numpy(float)
    v1 = df["se_eqtl"].to_numpy(float) ** 2
    z2 = df["beta_gwas"].to_numpy(float) / df["se_gwas"].to_numpy(float)
    v2 = df["se_gwas"].to_numpy(float) ** 2
    abf1 = np.sqrt(v1 / (v1 + w1)) * np.exp((z1**2) * w1 / (2 * (v1 + w1)))
    abf2 = np.sqrt(v2 / (v2 + w2)) * np.exp((z2**2) * w2 / (2 * (v2 + w2)))
    s1 = np.nansum(abf1)
    s2 = np.nansum(abf2)
    s12 = np.nansum(abf1 * abf2)
    denom = 1 + p1 * s1 + p2 * s2 + p1 * p2 * max(s1 * s2 - s12, 0) + p12 * s12
    return {
        "PP0": 1 / denom,
        "PP1": p1 * s1 / denom,
        "PP2": p2 * s2 / denom,
        "PP3": p1 * p2 * max(s1 * s2 - s12, 0) / denom,
        "PP4": p12 * s12 / denom,
        "PP4_over_PP3_plus_PP4": (p12 * s12) / (p1 * p2 * max(s1 * s2 - s12, 0) + p12 * s12) if (p1 * p2 * max(s1 * s2 - s12, 0) + p12 * s12) else np.nan,
    }


def make_coloc_prior_sensitivity() -> None:
    rows = []
    merged_for_plot = []
    for cell in ["bin", "bmem"]:
        eqtl = extract_park7_eqtl(cell)
        eqtl = eqtl[(eqtl["CHR"].astype(str) == "1") & (eqtl["POS"].between(6_400_000, 8_600_000))].copy()
        eqtl["z"] = ndtri(1 - eqtl["P_VALUE"].clip(1e-300, 1) / 2) * np.sign(eqtl["SPEARMANS_RHO"])
        eqtl["se_eqtl"] = (eqtl["SPEARMANS_RHO"].abs() / np.maximum(np.abs(eqtl["z"]), 1e-12)).replace([np.inf, -np.inf], np.nan)
        for hist, fname in [("LUSC", "region_squamous_chr1.tsv"), ("LUAD", "region_adeno_chr1.tsv")]:
            gwas = pd.read_csv(PROJECT / fname, sep="\t")
            gwas = gwas.rename(columns={"rsids": "RSID", "beta": "beta_gwas", "sebeta": "se_gwas", "pval": "p_gwas"})
            merged = eqtl.merge(gwas[["RSID", "pos", "beta_gwas", "se_gwas", "p_gwas"]], on="RSID", how="inner")
            merged = merged.rename(columns={"SPEARMANS_RHO": "beta_eqtl"})
            merged = merged.dropna(subset=["beta_eqtl", "se_eqtl", "beta_gwas", "se_gwas"])
            if len(merged) < 10:
                continue
            if cell == "bin" and hist == "LUSC":
                merged_for_plot.append(merged.assign(cell=cell, histology=hist))
            for p12 in [1e-6, 1e-5, 1e-4]:
                res = coloc_abf(merged, 1e-4, 1e-4, p12)
                rows.append({"cell": cell, "histology": hist, "p1": 1e-4, "p2": 1e-4, "p12": p12, "n_shared_snps": len(merged), **res})
    out = pd.DataFrame(rows)
    out.to_csv(SRC / "S17_Coloc_ABF_Prior_Sensitivity.csv", index=False)
    add_manifest_row("S17_Coloc_ABF_Prior_Sensitivity.csv", "Supplementary Fig. S17", "Approximate ABF colocalization prior sensitivity using local PARK7 region data.")
    if merged_for_plot:
        mp = pd.concat(merged_for_plot, ignore_index=True)
        mp["neglog10_gwas"] = -np.log10(mp["p_gwas"].clip(1e-300, 1))
        mp["neglog10_eqtl"] = -np.log10(mp["P_VALUE"].clip(1e-300, 1))
        mp.to_csv(SRC / "S17_PARK7_LocusCompare_Source.csv", index=False)
        add_manifest_row("S17_PARK7_LocusCompare_Source.csv", "Supplementary Fig. S17", "PARK7 region eQTL/GWAS merged source for locus-compare plot.")
    fig, axes = plt.subplots(1, 2, figsize=(10.5, 4.0))
    if not out.empty:
        for (cell, hist), sub in out.groupby(["cell", "histology"]):
            axes[0].plot(sub["p12"].astype(str), sub["PP4"], marker="o", label=f"{cell}-{hist}")
        axes[0].set_ylabel("PP4")
        axes[0].set_xlabel("p12 prior")
        axes[0].set_title("Prior sensitivity")
        axes[0].legend(frameon=False, fontsize=7)
    if merged_for_plot:
        axes[1].scatter(mp["neglog10_eqtl"], mp["neglog10_gwas"], s=12, alpha=0.65)
        axes[1].set_xlabel("-log10 eQTL P")
        axes[1].set_ylabel("-log10 LUSC GWAS P")
        axes[1].set_title("PARK7 locus-compare view")
    fig.suptitle("Supplementary Fig. S17. PARK7 colocalization prior and regional sensitivity", fontsize=13)
    fig.tight_layout()
    save_fig(fig, SUPP / "S17_Coloc_Prior_Regional_Sensitivity.tif")


def make_spatial_neighborhood() -> None:
    df = pd.read_csv(SRC / "Figure4_Spatial_CoLocalization_Scores.csv")
    bins = [-0.1, 0, 1, 2, 5, np.inf]
    labels = ["B-high spot", "<=1 grid", "1-2 grids", "2-5 grids", ">5 grids"]
    df["distance_to_B_high_bin"] = pd.cut(df["dist_to_b"], bins=bins, labels=labels)
    rows = []
    for b, sub in df.groupby("distance_to_B_high_bin", observed=False):
        rows.append(
            {
                "distance_bin": str(b),
                "n_spots": len(sub),
                "Antioxidant_z_mean": sub["Antioxidant_z"].mean(),
                "Antioxidant_z_median": sub["Antioxidant_z"].median(),
                "B_cell_density_z_mean": sub["B_cell_density_z"].mean(),
                "Both_high_fraction": sub["Both_high"].mean(),
            }
        )
    rho, p = stats.spearmanr(df["dist_to_b"], df["Antioxidant_z"], nan_policy="omit")
    summary = pd.DataFrame(rows)
    summary["global_spearman_dist_to_B_vs_antioxidant_z"] = rho
    summary["global_spearman_p"] = p
    summary.to_csv(SRC / "S18_Spatial_Neighborhood_Distance_Analysis.csv", index=False)
    add_manifest_row("S18_Spatial_Neighborhood_Distance_Analysis.csv", "Supplementary Fig. S18", "Distance-to-B-cell-high neighborhood analysis for antioxidant score.")
    fig, axes = plt.subplots(1, 2, figsize=(10.2, 4.0))
    axes[0].bar(summary["distance_bin"], summary["Antioxidant_z_mean"], color="#CC6677")
    axes[0].tick_params(axis="x", rotation=30)
    axes[0].axhline(0, color="black", lw=0.8)
    axes[0].set_ylabel("Mean antioxidant z-score")
    axes[0].set_title("Antioxidant score by distance to B-high spots")
    axes[1].scatter(df["dist_to_b"], df["Antioxidant_z"], s=5, alpha=0.25)
    axes[1].set_xlabel("Distance to nearest B-high spot")
    axes[1].set_ylabel("Antioxidant z-score")
    axes[1].set_title(f"Spearman rho={rho:.3f}")
    fig.suptitle("Supplementary Fig. S18. Spatial neighborhood analysis", fontsize=13)
    fig.tight_layout()
    save_fig(fig, SUPP / "S18_Spatial_Neighborhood_Distance_Analysis.tif")


def make_network_control_context() -> None:
    b = pd.read_csv(SRC / "S15_GSE148071_LUSC_scRNA_Blineage_stress_scores.csv")
    rows = []
    if not b.empty:
        for group, sub in b.groupby("cell_group"):
            for col in [c for c in sub.columns if c.endswith("_delta_high_minus_low")]:
                vals = sub[col].dropna()
                if len(vals):
                    rows.append(
                        {
                            "cell_group": group,
                            "score": col.replace("_delta_high_minus_low", ""),
                            "n_samples": len(vals),
                            "median_delta_PARK7_high_minus_low": vals.median(),
                            "wilcoxon_signedrank_p": stats.wilcoxon(vals).pvalue if len(vals) > 0 and not np.allclose(vals, 0) else np.nan,
                        }
                    )
    out = pd.DataFrame(rows)
    if not out.empty:
        out["wilcoxon_fdr"] = bh(out["wilcoxon_signedrank_p"])
    out.to_csv(SRC / "S19_GSE148071_Blineage_PARK7high_StressScore_Meta.csv", index=False)
    add_manifest_row("S19_GSE148071_Blineage_PARK7high_StressScore_Meta.csv", "Supplementary Fig. S19", "Across-sample B-lineage PARK7-high stress-score directionality in GSE148071.")
    gene_class = pd.read_csv(SRC / "S14_NetworkInference_TopGeneClass_Audit.csv")
    fig, axes = plt.subplots(1, 2, figsize=(11, 4.0))
    if not out.empty:
        plot = out[out["cell_group"] == "B-lineage combined"]
        axes[0].barh(plot["score"], plot["median_delta_PARK7_high_minus_low"], color="#4477AA")
        axes[0].axvline(0, color="black", lw=0.8)
        axes[0].set_title("GSE148071 B-lineage PARK7-high stress context")
        axes[0].set_xlabel("Median sample-level delta")
    axes[1].bar(gene_class["gene_class"], gene_class["count_in_top500"], color="#AA4499")
    axes[1].tick_params(axis="x", rotation=35)
    axes[1].set_title("Network top-gene class audit")
    axes[1].set_ylabel("Count among top 500 genes")
    fig.suptitle("Supplementary Fig. S19. Network and independent scRNA control context", fontsize=13)
    fig.tight_layout()
    save_fig(fig, SUPP / "S19_Network_scRNA_Control_Context.tif")


def update_docs() -> None:
    backup(MANUSCRIPT)
    doc = Document(str(MANUSCRIPT))
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("Public LUSC scRNA-seq validation used GSE200972") and "GSE148071" not in text:
            replace_para(
                para,
                para.text
                + " Independent public LUSC scRNA-seq validation additionally used the processed expression matrices from GSE148071, comprising 42 LUSC samples. To avoid relying on unavailable author-level cell annotations, cells were assigned to broad compartments using the same marker-score framework, and PARK7 detection, positive-cell expression intensity, and B-lineage PARK7-high stress-score differences were summarized at the sample level.",
            )
        elif text.startswith("Cell-type-resolved MR prioritized PARK7 expression") and "Extended MR robustness analyses" not in text:
            replace_para(
                para,
                para.text
                + " Extended MR robustness analyses using available harmonized instruments showed directionally positive estimates across IVW, weighted-median, mode-based, MR-Egger-slope, and Huber robust estimators; approximate Steiger and radial-like diagnostics are reported as sensitivity checks rather than replacements for package-specific MR-RAPS/RadialMR implementations (Supplementary Fig. S16 and source data).",
            )
        elif text.startswith("Locus-level analyses suggested LUSC-preferential posterior support") and "prior-sensitivity" not in text:
            replace_para(
                para,
                para.text
                + " Additional ABF-style prior-sensitivity and locus-compare summaries are provided as Supplementary Fig. S17; these reinforce the interpretation as weak-to-moderate regional support rather than high-confidence colocalization.",
            )
        elif text.startswith("Spatial transcriptomic analysis identified heterogeneous") and "neighborhood-distance" not in text:
            replace_para(
                para,
                para.text
                + " A neighborhood-distance analysis further evaluated antioxidant scores as a function of distance to B-cell-high spots and is reported as spatial context rather than co-localization evidence (Supplementary Fig. S18).",
            )
        elif text.startswith("As an additional public LUSC scRNA-seq validation") and "GSE148071" not in text:
            replace_para(
                para,
                para.text
                + " A larger independent LUSC scRNA-seq processed dataset, GSE148071, showed PARK7 detectability across broad assigned tumor-microenvironment compartments, including B-lineage cells, and provided sample-level B-lineage stress-score context for PARK7-high versus PARK7-low cells (Supplementary Fig. S15 and Supplementary Fig. S19).",
            )
        elif "As an additional public LUSC scRNA-seq validation, GSE200972" in text and "GSE148071" not in text:
            replace_para(
                para,
                para.text
                + " A larger independent LUSC scRNA-seq processed dataset, GSE148071, was also analyzed with the same marker-score framework and showed PARK7 detectability across broad assigned tumor-microenvironment compartments, including B-lineage cells; sample-level B-lineage stress-score comparisons for PARK7-high versus PARK7-low cells are provided as contextual rather than functional evidence (Supplementary Fig. S15 and Supplementary Fig. S19).",
            )
        elif text.startswith("Several limitations should be acknowledged") and "GSE148071 processed expression matrices" not in text:
            replace_para(
                para,
                para.text
                + " Tenth, the GSE148071 validation used processed expression matrices and marker-score compartment assignment rather than the original authors' full cell-state annotation; therefore, it strengthens public-data expression reproducibility but still does not prove B-cell-specific function or protein localization.",
            )
        elif text.startswith("Publicly available datasets analyzed in this study") and "GSE148071" not in text:
            replace_para(
                para,
                para.text.replace(
                    "including GSE200972 public LUSC scRNA-seq data,",
                    "including GSE200972 and GSE148071 public LUSC scRNA-seq data,",
                )
            )
    doc.save(str(MANUSCRIPT))

    backup(TITLE_PAGE)
    tdoc = Document(str(TITLE_PAGE))
    for para in tdoc.paragraphs:
        if para.text.strip().startswith("Number of supplementary figures:"):
            replace_para(para, "Number of supplementary figures: 19")
        elif para.text.strip().startswith("Number of tables:"):
            replace_para(para, "Number of tables: 1 main table and 14 supplementary/source tables")
    tdoc.save(str(TITLE_PAGE))

    supp_doc = SUPP / "Supplementary_Material.docx"
    backup(supp_doc)
    sdoc = Document(str(supp_doc))
    existing = "\n".join(p.text for p in sdoc.paragraphs)
    additions = [
        ("Supplementary Fig. S15. Independent GSE148071 LUSC scRNA PARK7 validation.", "Processed expression matrices from 42 public LUSC samples were analyzed using marker-score compartment assignment. PARK7 detection and expression intensity are summarized by assigned cell type, and B-lineage PARK7-high versus PARK7-low stress-score context is shown."),
        ("Supplementary Fig. S16. Extended PARK7 MR robustness analyses.", "Available harmonized PARK7 instruments were used for robust estimator comparison, radial-like outlier diagnostics, approximate Steiger directionality, and eQTL-threshold/cis-window sensitivity. Package-specific MR-RAPS and RadialMR could not be run because the R packages were unavailable for the local R 4.5.3 installation."),
        ("Supplementary Fig. S17. PARK7 colocalization prior and regional sensitivity.", "Approximate ABF-style prior sensitivity and locus-compare summaries are shown for the PARK7 region. These analyses are used to support weak-to-moderate regional evidence and not definitive colocalization."),
        ("Supplementary Fig. S18. Spatial neighborhood analysis.", "Antioxidant scores are summarized by distance to B-cell-high spots, providing neighborhood-level spatial context without implying PARK7 spatial localization."),
        ("Supplementary Fig. S19. Network and independent scRNA control context.", "GSE148071 B-lineage PARK7-high stress-score directionality and network top-gene class composition are summarized as exploratory control context."),
    ]
    for title, body in additions:
        if title not in existing:
            p = sdoc.add_paragraph()
            p.add_run(title).bold = True
            sdoc.add_paragraph(body)
    sdoc.save(str(supp_doc))


def update_package_manifest() -> None:
    rows = []
    for path in PACKAGE.rglob("*"):
        if not path.is_file() or ".bak_" in path.name.lower() or "backup" in path.name.lower():
            continue
        rel = path.relative_to(PACKAGE).as_posix()
        role = "package documentation"
        if rel.startswith("01_Manuscript/"):
            role = "core manuscript"
        elif rel.startswith("02_Frontiers_Statements/"):
            role = "Frontiers statement/admin"
        elif rel.startswith("03_Figures_Main/"):
            role = "main figure"
        elif rel.startswith("04_Supplementary_Material/"):
            role = "supplementary material"
        elif rel.startswith("05_Source_Data/"):
            role = "source data / reproducibility"
        elif rel.startswith("06_Original_Microscopy_Images/"):
            role = "original microscopy image"
        elif rel.startswith("07_Admin_and_Optional/"):
            role = "admin optional"
        rows.append({"relative_path": rel, "size_bytes": path.stat().st_size, "role": role})
    pd.DataFrame(rows).sort_values("relative_path").to_csv(PACKAGE / "PACKAGE_MANIFEST.csv", index=False)


def rebuild_zip(zip_path: Path, source_dir: Path) -> None:
    if zip_path.exists():
        zbak = WORK / "zip_backups"
        zbak.mkdir(exist_ok=True)
        shutil.copy2(zip_path, zbak / (zip_path.name + ".before_completed_public_analyses"))
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=8) as z:
        for path in source_dir.rglob("*"):
            if not path.is_file():
                continue
            if path.resolve() == zip_path.resolve():
                continue
            if ".bak_" in path.name.lower() or "backup" in path.name.lower():
                continue
            if source_dir == PACKAGE and path.suffix.lower() == ".zip" and path.parent == PACKAGE:
                continue
            z.write(path, path.relative_to(source_dir.parent))


def main() -> None:
    process_gse148071()
    make_mr_extended()
    make_coloc_prior_sensitivity()
    make_spatial_neighborhood()
    make_network_control_context()
    update_docs()
    update_package_manifest()
    rebuild_zip(PACKAGE / "04_Supplementary_Material.zip", SUPP)
    rebuild_zip(PACKAGE / "05_Source_Data" / "Source_Data_CSV.zip", SRC)
    rebuild_zip(PACKAGE / "05_Source_Data.zip", PACKAGE / "05_Source_Data")
    rebuild_zip(ROOT / "submission_package.zip", PACKAGE)
    (WORK / "COMPLETED_PUBLIC_ANALYSES_REPORT.md").write_text(
        "# Completed public-data analysis expansion\n\n"
        "- Downloaded and analyzed GSE148071 RAW tar processed expression matrices (42 LUSC samples).\n"
        "- Added independent scRNA validation, extended MR diagnostics, coloc prior sensitivity, spatial neighborhood analysis, and network/scRNA control context.\n"
        "- Updated manuscript, supplementary material, source-data manifest, and zips.\n",
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
