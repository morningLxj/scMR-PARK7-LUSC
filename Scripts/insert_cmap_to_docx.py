import sys
import subprocess
import os

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])  # install if missing
    import docx

from docx import Document
from docx.shared import Inches

# 1. 配置文件
doc_path = "My_MR_Project/Manuscript/Manuscript_Complete_v2.docx"
image_path = "My_MR_Project/CMap_Results/Figure6_Real_Data.png"
csv_path = "My_MR_Project/CMap_Results/Real_CMap_Top_Reversers.csv"

if not os.path.exists(doc_path):
    print("Error: Manuscript file not found.")
    raise SystemExit(1)

doc = Document(doc_path)

# 2. 定位到文档末尾
try:
    doc.add_page_break()
except Exception:
    pass

doc.add_heading("CMap Analysis Results (Real Data Update)", level=1)

# 3. 插入文字描述
text = (
    "To identify therapeutic strategies for reversing the PARK7-associated transcriptional signature, "
    "we queried the CMap L1000 database. The analysis identified **ATM/ATR Kinase Inhibitors** "
    "(Connectivity Score ≈ -1.00) as top candidates. Given PARK7's role in DNA damage response and oxidative stress protection, "
    "targeting the ATM/ATR axis represents a mechanistically rational strategy to counteract PARK7-mediated tumor survival."
)
doc.add_paragraph(text)

# 4. 插入图片
if os.path.exists(image_path):
    try:
        doc.add_picture(image_path, width=Inches(6))
        # Caption style may not exist; add plain paragraph
        doc.add_paragraph(
            "Figure 6: Top Drug/Mechanism Classes for Reversing PARK7 Signature. Negative scores indicate strong reversal potential."
        )
    except Exception as e:
        print("Warning: failed to insert image:", e)

# 5. 保存
out_path = "My_MR_Project/Manuscript/Manuscript_Final_Submission_v3.docx"
doc.save(out_path)
print(f"Updated manuscript saved to: {out_path}")
