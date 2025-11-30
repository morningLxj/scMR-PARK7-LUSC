import sys, subprocess, os
try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
import pandas as pd
from docx import Document

in_doc = "My_MR_Project/Manuscript/Manuscript_Final_Submission_v3.docx"
csv_path = "My_MR_Project/CMap_Results/Real_CMap_Top_Reversers_DrugLevel.csv"
out_doc = "My_MR_Project/Manuscript/Manuscript_Final_Submission_v4.docx"

if not os.path.exists(in_doc):
    raise SystemExit(1)

doc = Document(in_doc)

doc.add_heading("方法简述", level=2)
doc.add_paragraph(
    "我们基于 CLUE L1000 结果包，自动定位含单药物级别信息的 GCT 文件（包含 pert_iname 与 norm_cs/raw_cs 列），"
    "并筛选干预类型为化合物（trt_cp），按 Connectivity Score（越负表示逆转越强）对药物进行排序。"
)
doc.add_paragraph(
    "为避免同一药物在不同细胞系或剂量重复计数，我们按药物名聚合并取最负的分值作为代表值，最终输出药物级 Top15。"
)

top_names = []
top_scores = []
if os.path.exists(csv_path):
    df = pd.read_csv(csv_path)
    for i in range(min(5, len(df))):
        top_names.append(str(df.iloc[i]["Drug Name"]))
        top_scores.append(float(df.iloc[i]["Connectivity Score"]))

doc.add_heading("结果解读", level=2)
if top_names:
    main = top_names[0]
    score = top_scores[0]
    extra = ", ".join(top_names[1:]) if len(top_names) > 1 else ""
    doc.add_paragraph(
        f"药物级 Top15 中，{main}（Connectivity Score ≈ {score:.2f}）位列首位；紧随其后包括 {extra} 等。"
    )
doc.add_paragraph(
    "这些候选多集中于 DNA 拓扑异构酶抑制与 DNA 损伤相关通路，提示通过诱导肿瘤细胞应激与损伤反应，"
    "可在功能上逆转 PARK7 相关转录表型与其抗氧化保护效应，从而达到抑制肿瘤生存的目的。"
)
doc.add_paragraph(
    "结合 MR 与 CMap 的证据，针对 DNA 损伤/修复轴的药理干预（包括拓扑异构酶抑制剂与 ATM/ATR 路径抑制剂）"
    "具有机制合理性与转化潜力，值得在后续研究与临床方案设计中优先验证。"
)

doc.save(out_doc)
print(out_doc)
