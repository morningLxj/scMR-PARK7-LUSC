import sys, subprocess, os
try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])  
    import docx
import pandas as pd
from docx import Document
from docx.shared import Inches

in_doc = "My_MR_Project/Manuscript/Manuscript_Final_Submission_v4.docx"
out_doc = "My_MR_Project/Manuscript/Manuscript_Final_Submission_v5.docx"
drug_csv = "My_MR_Project/CMap_Results/Real_CMap_Top_Reversers_DrugLevel.csv"
moa_csv = "My_MR_Project/CMap_Results/Real_CMap_Top_Reversers_MOA.csv"
flow_img = "My_MR_Project/CMap_Results/Method_Flowchart.png"
supp_bar = "My_MR_Project/CMap_Results/Supplement/Drug_Top15_Bar.png"
supp_violin = "My_MR_Project/CMap_Results/Supplement/Score_Violin_CellLine.png"

if not os.path.exists(in_doc):
    raise SystemExit(1)

doc = Document(in_doc)

doc.add_heading("并列表格：药物级与机制级 Top15", level=2)

if os.path.exists(drug_csv):
    df = pd.read_csv(drug_csv)
    df = df.head(15)
    table = doc.add_table(rows=df.shape[0]+1, cols=2)
    table.style = 'Light Shading'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Drug Name'
    hdr_cells[1].text = 'Connectivity Score'
    for i in range(df.shape[0]):
        row_cells = table.rows[i+1].cells
        row_cells[0].text = str(df.iloc[i]['Drug Name'])
        row_cells[1].text = f"{float(df.iloc[i]['Connectivity Score']):.4f}"

if os.path.exists(moa_csv):
    dfm = pd.read_csv(moa_csv)
    dfm = dfm.head(15)
    doc.add_paragraph("")
    table2 = doc.add_table(rows=dfm.shape[0]+1, cols=3)
    table2.style = 'Light Shading'
    h = table2.rows[0].cells
    h[0].text = 'Mechanism'
    h[1].text = 'Connectivity Score'
    h[2].text = 'Example Drug'
    for i in range(dfm.shape[0]):
        r = table2.rows[i+1].cells
        r[0].text = str(dfm.iloc[i]['Mechanism'])
        r[1].text = f"{float(dfm.iloc[i]['Connectivity Score']):.4f}"
        r[2].text = str(dfm.iloc[i].get('Example Drug',''))

if os.path.exists(flow_img):
    doc.add_paragraph("")
    doc.add_heading("方法流程图", level=2)
    doc.add_picture(flow_img, width=Inches(6))

doc.add_heading("Supplement: 额外图示", level=2)
if os.path.exists(supp_bar):
    doc.add_picture(supp_bar, width=Inches(6))
if os.path.exists(supp_violin):
    doc.add_picture(supp_violin, width=Inches(6))

doc.save(out_doc)
print(out_doc)
