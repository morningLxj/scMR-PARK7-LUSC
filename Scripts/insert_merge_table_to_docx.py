import sys, subprocess, os
try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])  
    import docx
import pandas as pd
from docx import Document
from docx.shared import Inches

in_doc = "My_MR_Project/Manuscript/Manuscript_Final_Submission_v6.docx"
out_doc = "My_MR_Project/Manuscript/Manuscript_Final_Submission_v7.docx"
csv_path = "My_MR_Project/CMap_Results/DrugMechanism_Top15.csv"
bar_img = "My_MR_Project/CMap_Results/DrugMechanism_Top15_Bar.png"
pc3_csv = "My_MR_Project/CMap_Results/DrugLevel_Top15_PC3.csv"
hek_csv = "My_MR_Project/CMap_Results/DrugLevel_Top15_HEK293.csv"
violin_img = "My_MR_Project/CMap_Results/Supplement/Score_Violin_CellLine.png"
mech_imgs_dir = "My_MR_Project/CMap_Results/Supplement"

if not os.path.exists(in_doc):
    raise SystemExit(1)

doc = Document(in_doc)

doc.add_heading("药物-机制合并表格", level=2)
if os.path.exists(csv_path):
    df = pd.read_csv(csv_path)
    df = df.head(15)
    table = doc.add_table(rows=df.shape[0]+1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Drug Name"
    hdr[1].text = "Mechanism"
    hdr[2].text = "Connectivity Score"
    for i in range(df.shape[0]):
        r = table.rows[i+1].cells
        r[0].text = str(df.iloc[i]["Drug Name"])
        r[1].text = str(df.iloc[i]["Mechanism"])
        r[2].text = f"{float(df.iloc[i]['Connectivity Score']):.4f}"

if os.path.exists(bar_img):
    doc.add_paragraph("")
    doc.add_picture(bar_img, width=Inches(6))

# stratified tables
doc.add_heading("Supplement: PC3/HEK293 分层 Top15 表", level=2)
def insert_two_col_table(path):
    if os.path.exists(path):
        df = pd.read_csv(path)
        df = df.head(15)
        table = doc.add_table(rows=df.shape[0]+1, cols=2)
        hdr = table.rows[0].cells
        hdr[0].text = "Drug Name"
        hdr[1].text = "Connectivity Score"
        for i in range(df.shape[0]):
            r = table.rows[i+1].cells
            r[0].text = str(df.iloc[i]["Drug Name"])
            r[1].text = f"{float(df.iloc[i]['Connectivity Score']):.4f}"
        doc.add_paragraph("")
insert_two_col_table(pc3_csv)
insert_two_col_table(hek_csv)

if os.path.exists(violin_img):
    doc.add_picture(violin_img, width=Inches(6))

# embed first 6 mechanism-specific charts
doc.add_heading("Supplement: 机制特异条形图（Top6）", level=2)
if os.path.isdir(mech_imgs_dir):
    imgs = [os.path.join(mech_imgs_dir, f) for f in os.listdir(mech_imgs_dir) if f.startswith("Mechanism_") and f.endswith("_Bar.png")]
    imgs.sort()
    for p in imgs[:6]:
        doc.add_picture(p, width=Inches(5.5))

doc.save(out_doc)
print(out_doc)
